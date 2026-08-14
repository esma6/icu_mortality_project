from __future__ import annotations

from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
GEN = ROOT / "paper_revision" / "generated"
OUT = ROOT / "paper_revision" / "MIMICIII_Spark_ETL_Makale_Literatur_Temelli_v8_final.docx"

SCENARIO_LABEL = {
    "1-node": "local[2]",
    "1-node-4c": "local[4]",
    "1-node-8c": "local[8]",
    "2-node": "2 düğüm",
    "3-node": "3 düğüm",
}


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_cell_text(cell, text: str, bold: bool = False, color: str | None = None) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(str(text))
    r.bold = bold
    r.font.name = "Arial"
    r.font.size = Pt(8.2)
    if color:
        r.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    hdr = table.rows[0]
    set_repeat_table_header(hdr)
    for i, h in enumerate(headers):
        set_cell_text(hdr.cells[i], h, bold=True, color="FFFFFF")
        set_cell_shading(hdr.cells[i], "1F4E79")
        if widths:
            hdr.cells[i].width = Cm(widths[i])
    for ridx, row in enumerate(rows):
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], value)
            if ridx % 2:
                set_cell_shading(cells[i], "EAF2F8")
            if widths:
                cells[i].width = Cm(widths[i])
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    return table


def add_caption(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="Caption")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(text)


def add_figure(doc: Document, filename: str, caption: str, width_cm: float = 16.7) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.keep_with_next = True
    shape = p.add_run().add_picture(str(GEN / filename), width=Cm(width_cm))
    shape._inline.docPr.set("descr", caption)
    shape._inline.docPr.set("title", filename)
    add_caption(doc, caption)


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    p = doc.add_heading(text, level=level)
    p.paragraph_format.keep_with_next = True


def add_body(doc: Document, text: str, bold_lead: str | None = None) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0.65)
    if bold_lead and text.startswith(bold_lead):
        p.add_run(bold_lead).bold = True
        p.add_run(text[len(bold_lead):])
    else:
        p.add_run(text)


def add_bullet(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Cm(0.8)
    p.paragraph_format.first_line_indent = Cm(-0.4)
    p.add_run(text)


def fmt_pm(mean, sd):
    return f"{mean:.1f} ± {sd:.1f}"


def runtime_rows(summary: pd.DataFrame, workload: str):
    d = summary[summary.workload == workload]
    rows = []
    for _, r in d.iterrows():
        rows.append([
            SCENARIO_LABEL[r.scenario],
            fmt_pm(r.extract_seconds_mean, r.extract_seconds_std),
            fmt_pm(r.transform_seconds_mean, r.transform_seconds_std),
            fmt_pm(r.load_seconds_mean, r.load_seconds_std),
            fmt_pm(r.total_seconds_mean, r.total_seconds_std),
            f"{r.total_seconds_median:.1f} [{r.total_seconds_q1:.1f}–{r.total_seconds_q3:.1f}]",
        ])
    return rows


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.25)
    section.bottom_margin = Cm(2.25)
    section.left_margin = Cm(2.35)
    section.right_margin = Cm(2.35)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(10.5)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.line_spacing = 1.08
    normal.paragraph_format.space_after = Pt(5)

    for name, size, color in [("Title", 17, "17365D"), ("Heading 1", 13, "17365D"), ("Heading 2", 11.5, "1F4E79"), ("Heading 3", 10.5, "1F4E79")]:
        st = styles[name]
        st.font.name = "Arial"
        st.font.size = Pt(size)
        st.font.bold = True
        st.font.color.rgb = RGBColor.from_string(color)
        st.paragraph_format.space_before = Pt(8)
        st.paragraph_format.space_after = Pt(4)
        st.paragraph_format.keep_with_next = True

    cap = styles["Caption"]
    cap.font.name = "Arial"
    cap.font.size = Pt(8.5)
    cap.font.italic = True
    cap.font.color.rgb = RGBColor.from_string("404040")
    cap.paragraph_format.space_after = Pt(7)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("MIMIC-III Spark ETL — hakem revizyonlu makale")
    run.font.name = "Arial"
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor.from_string("777777")
    run.add_tab()
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    footer._p.append(fld)


def build() -> None:
    summary = pd.read_csv(GEN / "runtime_summary.csv")
    omnibus = pd.read_csv(GEN / "omnibus_tests.csv")
    pairwise = pd.read_csv(GEN / "pairwise_tests_local8.csv")
    cpu = pd.read_csv(GEN / "cpu_summary.csv")

    doc = Document()
    configure_document(doc)

    title = doc.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.add_run("Kaynak Kısıtlı Tek-Host Ortamında Klinik Spark ETL:\nMIMIC-III Üzerinde İş Yükü Granülerliği ve Yürütme Topolojisinin Tekrarlı Deneysel Değerlendirmesi")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("[Yazar adı]  |  [Kurum / bölüm]  |  [E-posta]").italic = True
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("Araştırma makalesi — bağımsız yöntem ve istatistik hakemliği sonrası revizyon").bold = True

    add_heading(doc, "Öz", 1)
    add_body(doc, "Arka plan: Klinik EHR veri hatları 2020 sonrasında standartlaştırılmış zaman serilerinden olay-merkezli ve kurumlar arası taşınabilir temsillere ilerlemiştir; buna karşın kaynak kısıtlı ortamlarda özellik granülerliği ile Spark yürütme topolojisinin birlikte etkisi yeterince ölçülmemiştir. Amaç: Aynı fiziksel kaynak bütçesinde Spark local iş parçacığı paralelliği ile eş-konumlu standalone worker topolojilerini iki MIMIC-III özellik rejiminde karşılaştırmak. Yöntem: Docker Compose ile kapsayıcılaştırılmış Spark 3.5.1 hattı, 12 mantıksal işlemci ve tek NVMe SSD'li aynı Windows 11 hostunda local[2], local[4], local[8], master+bir worker ve master+iki worker olarak çalıştırıldı. Kabul düzeyindeki compact ürün 58.976, altı saatlik vital pencereli timeseries ürün 1.180.395 satır üretti. Her iş yükü–senaryo çifti altı kez çalıştırıldı; r01 warmup olarak dışlandı ve n=5 ölçüm Welch ANOVA, Holm-düzeltilmiş Welch karşılaştırmaları ve host telemetrisiyle analiz edildi. Bulgular: En düşük ortalama süre compact için local[8]'de 204,8±2,9 s, timeseries için 263,9±30,1 s idi. Bir-worker standalone süreleri 226,8±5,1 ve 291,7±24,1 s; iki-worker süreleri 320,7±30,5 ve 476,8±133,0 s idi. Welch omnibus testleri compact (p<0,001) ve timeseries (p=0,027) için anlamlıydı; ancak timeseries local[8] ikili farkları Holm düzeltmesinden sonra anlamlı değildi. Sonuç: İncelenen tek-host ortamında worker kapsayıcısı eklemek fiziksel kapasite eklememiş ve local[8]'i geçmemiştir. Bulgular gerçek çok-hostlu strong scaling kanıtı değil, sabit kaynak bütçesinde yürütme-topolojisi karakterizasyonudur; sonuç, Quick-MIMIC'in azalan paralel verimlilik bulgusuyla nitel olarak uyumludur.")
    p = doc.add_paragraph()
    p.add_run("Anahtar sözcükler: ").bold = True
    p.add_run("MIMIC-III; Apache Spark; ETL; zaman serisi; Docker; performans karakterizasyonu; yoğun bakım verisi")

    add_heading(doc, "Abstract", 1)
    add_body(doc, "Background: Since 2020, clinical EHR pipelines have progressed from standardized time-series extraction toward event-centric and interoperable representations, yet the joint effect of feature granularity and Spark execution topology remains under-characterized in resource-constrained environments. Objective: To compare local thread parallelism with co-located standalone Spark workers under a fixed physical resource budget. Methods: A Docker Compose–containerized Spark 3.5.1 ETL pipeline processed MIMIC-III v1.4 on one Windows 11 host with 12 logical processors and one NVMe SSD. Five configurations—local[2], local[4], local[8], master plus one worker, and master plus two workers—were tested on an admission-level compact workload (58,976 rows) and a six-hour vital-window workload (1,180,395 rows). Six runs were performed per workload–configuration pair; r01 was excluded as warm-up and five measured runs were analyzed using Welch ANOVA, Holm-adjusted Welch contrasts, and host telemetry. Results: Local[8] had the lowest mean runtime for compact (204.8±2.9 s) and timeseries (263.9±30.1 s). One-worker standalone runtimes were 226.8±5.1 and 291.7±24.1 s; two-worker runtimes were 320.7±30.5 and 476.8±133.0 s. Welch omnibus tests were significant for compact (p<0.001) and timeseries (p=0.027), but no timeseries pairwise contrast against local[8] remained significant after Holm correction. Conclusion: On this single host, adding worker containers did not add physical capacity and did not outperform local[8]. This is a fixed-budget topology characterization, not evidence of multi-host strong scaling; the diminishing efficiency is qualitatively consistent with Quick-MIMIC.")
    p = doc.add_paragraph()
    p.add_run("Keywords: ").bold = True
    p.add_run("MIMIC-III; Apache Spark; ETL; time series; Docker; performance characterization; critical care data")

    add_heading(doc, "1. Giriş", 1)
    add_body(doc, "Elektronik sağlık kayıtları (EHR), klinik bakım sırasında farklı amaçlarla ve düzensiz zaman aralıklarında üretilen laboratuvar, vital bulgu, ilaç, tanı, işlem ve metin olaylarını bir araya getirir. Bu veriler araştırma için yüksek değer taşısa da operasyonel veri tabanları çoğunlukla analitik yorumlanabilirlik yerine kayıt ve hizmet sunumu için tasarlanmıştır. MIMIC-III ve onu izleyen MIMIC-IV gibi açık fakat kontrollü erişimli kaynaklar, yoğun bakım verisi üzerinde karşılaştırılabilir araştırmayı mümkün kılmış; aynı zamanda ilişkisel tablolardan kohort ve zaman-serisi üretmenin karmaşıklığını görünür hâle getirmiştir [1,2,16].")
    add_body(doc, "Veri hazırlama klinik makine öğrenmesinden önceki nötr bir mühendislik adımı değildir. Birim dönüşümü, eşdeğer itemid'lerin birleştirilmesi, zaman referansının seçimi, boş pencerelerin temsili, aykırı değer filtresi ve tahmin anından sonraki olayların dışlanması; hem örneklem büyüklüğünü hem de klinik sonucun görünür ilişkilerini değiştirir. Sağlıkta ML çalışmalarının kod ve veri erişilebilirliği bakımından diğer ML alanlarının gerisinde kaldığını gösteren incelemeler, yeniden üretilebilir veri hazırlama tanımını bilimsel geçerliliğin temel bileşeni olarak konumlandırmıştır [17].")
    add_body(doc, "2020–2022 dönemindeki MIMIC-Extract, FIDDLE, Clairvoyance, COP-E-CAT ve MIMIC-IV Data Pipeline çalışmaları bu probleme farklı yanıtlar vermiştir. MIMIC-Extract klinik olarak eşdeğer değişkenleri birleştirerek saatlik ICU zaman serileri üretmiş; FIDDLE veri kümesinden bağımsız ön filtre–dönüşüm–son filtre yapısı geliştirmiş; Clairvoyance klinik zaman serisi modellemesini modülerleştirmiş; COP-E-CAT ve MIMIC-IV Data Pipeline ise kohort ve görev yapılandırmasını MIMIC-IV'e taşımıştır [3,4,18–20]. Ortak yönleri, ham EHR ile model arasında açık ve tekrar kullanılabilir bir temsil katmanı kurmalarıdır.")
    add_body(doc, "2023 sonrasında araştırma gündemi yalnız sabit zaman pencereli özellik üretiminden olay-merkezli ve kurumlar arası taşınabilir veri katmanlarına genişlemiştir. EventStreamGPT sürekli-zaman karmaşık olay akışlarını foundation-model ölçeğinde işlemek için yapılandırma güdümlü bir kütüphane sunmuş; EHRSHOT, FEMR üzerinden hasta düzeyi global veri bölümü ve 15 görevli bir benchmark oluşturmuş; FHIR tabanlı bildirimsel hatlar klinik bilgi ile veri mühendisliği bilgisini ayırmayı hedeflemiştir [21–23]. MEDS yaklaşımı 2024'te başlayan çalışmaları 2026'da ortak, minimal ve event-centric bir Health AI veri standardı olarak birleştirmiştir [24]. Bu gelişim, yalnız hızlı ETL değil; semantik açıklık, provenance, leakage kontrolü ve algoritma taşınabilirliği beklentisini yükseltmiştir.")
    add_body(doc, "Bununla birlikte bu klinik veri hazırlama literatürünün büyük bölümü yürütme topolojisini deneysel değişken olarak ele almaz. Spark gibi dağıtık motorlarda çalışma süresi; veri büyüklüğü yanında seri kalan iş, shuffle, görev başlatma, JVM/serileştirme, bölümleme ve depolama paylaşımına bağlıdır [13,14]. `local[K]` aynı JVM içinde K worker iş parçacığı kullanırken standalone mod driver ve executor süreçlerini Spark master üzerinden koordine eder [5,6]. Bağımsız fiziksel CPU, bellek, ağ ve disk eklenmediğinde daha fazla worker kapsayıcısı gerçek cluster ölçeklenmesi sağlamaz; paylaşılan cache, bellek veri yolu ve I/O üzerinde çekişme yaratabilir [25,26].")
    add_body(doc, "MIMIC'e özgü en yakın performans çalışması Quick-MIMIC'tir. Dou ve arkadaşları MIMIC-III ve MIMIC-IV için çok süreçli ve MPI tabanlı multimodal çıkarımı değerlendirmiş; MIMIC-III'de dört süreçle 1,50× hızlanma fakat 0,37 paralel verimlilik ve artan I/O sınırlaması bildirmiştir [27]. Bu çalışma, paralellik ile veri çıkarımı arasındaki ilişkiyi doğrudan incelemesi bakımından önemli bir ilerlemedir. Ancak Spark'ın local iş parçacığı ve aynı fiziksel hostta eş-konumlu standalone worker topolojilerini, klinik özellik granülerliği değişirken tekrarlı koşular ve host telemetrisiyle karşılaştırmamıştır.")
    add_body(doc, "Bu boşluk özellikle eğitim, prototipleme ve kaynak kısıtlı araştırma ortamları için önemlidir. Bu ortamlarda Docker Compose ile bir master ve birden fazla worker oluşturmak sıklıkla 'düğüm ekleme' olarak yorumlanmaktadır; oysa süreç sayısının artması fiziksel kapasiteyi artırmaz. Yanlış yorumlanan hızlanma sonuçları, hem sistem seçimini hem de bilimsel ölçeklenebilirlik iddiasını bozabilir. Dolayısıyla araştırma problemi, Spark'ın genel olarak ölçeklenip ölçeklenmediği değil; aynı kaynak bütçesinde yürütme modelinin ve klinik özellik granülerliğinin uçtan uca ETL süresini nasıl değiştirdiğidir.")

    add_heading(doc, "1.1. Amaç ve araştırma soruları", 2)
    add_body(doc, "Çalışmanın amacı, MIMIC-III tabanlı iki özellik üretim rejiminde aynı fiziksel host üzerindeki Spark yürütme yapılandırmalarını, çıktı eşdeğerliği korunarak ve tekrarlı ölçümlerle karakterize etmektir. İnceleme aşağıdaki araştırma sorularına dayanır:")
    add_bullet(doc, "AS1: Kabul-seviyeli compact temsil ile vital bulguların altı saatlik pencerelere ayrıldığı daha yüksek granülerlikli temsil arasında uçtan uca süre ve aşama bileşimi nasıl değişmektedir?")
    add_bullet(doc, "AS2: Aynı kaynak bütçesinde local[2]/local[4]/local[8] iş parçacığı paralelliği ile bir- ve iki-worker Spark standalone topolojileri arasında hangi süre ve değişkenlik farkları gözlenmektedir?")
    add_bullet(doc, "AS3: Host CPU ve bellek telemetrisi, aynı hostta worker eklenmesinin gözlenen performansını ne ölçüde açıklamaktadır?")
    add_bullet(doc, "AS4: Bulgular, Quick-MIMIC'in çok süreçli/MPI paralellik sonuçları ve eş-konumlu kaynak çekişmesi literatürüyle hangi noktalarda örtüşmekte veya ayrışmaktadır?")

    add_heading(doc, "1.2. Katkılar", 2)
    add_body(doc, "Bu çalışma yeni bir klinik veri standardı veya klinik tahmin modeli önermemektedir. Katkıları daha dar ve sınanabilir dört noktadadır:")
    add_bullet(doc, "Aynı MIMIC-III Parquet girdileri üzerinde kabul-seviyeli ve vital-pencere-seviyeli iki iş yükünü tek bir Spark ETL hattında karşılaştırması;")
    add_bullet(doc, "fiziksel kaynak eklemeden local ve standalone yürütme topolojilerini ayırarak 'kapsayıcı sayısı = cluster ölçeği' varsayımını deneysel olarak sınaması;")
    add_bullet(doc, "her senaryo için warmup sonrası beş koşu, ham noktalar, heteroskedastisiteye dayanıklı testler ve host CPU/bellek telemetrisi raporlaması;")
    add_bullet(doc, "negatif ölçeklenme ve istatistiksel belirsizliği saklamadan, sonucu Quick-MIMIC ve güncel EHR veri katmanı literatürü içinde konumlandırması.")

    add_heading(doc, "2. Arka Plan ve İlgili Çalışmalar", 1)
    add_heading(doc, "2.1. Klinik olay verisinin zamansal temsili", 2)
    add_body(doc, "Klinik olay verisi doğal olarak düzensiz örneklenir: ölçüm sıklığı hastanın durumu ve bakım kararlarıyla ilişkilidir. Sabit saatlik veya altı saatlik agregasyon, farklı frekanstaki olayları ortak bir tensöre dönüştürmeyi kolaylaştırırken ölçüm sıklığı bilgisini azaltabilir; boş pencere üretimi, imputasyon ve maskeleme tercihleri modele ek varsayımlar taşır. MIMIC-Extract ve HiRID-ICU-Benchmark, görev ve etiket tanımlarıyla birlikte yeniden üretilebilir zaman-serisi üretiminin karşılaştırılabilirlik için önemini göstermiştir [3,28]. EventStreamGPT ve MEDS ise olay zamanını koruyan seyrek/event-centric temsillere yönelerek sabit grid zorunluluğunu azaltmıştır [21,24].")

    add_heading(doc, "2.2. EHR veri hazırlama hatlarının evrimi", 2)
    add_body(doc, "MIMIC-Extract ve FIDDLE'ın 2020'de belirginleştirdiği ilk kuşak problem, dağınık ilişkisel tablolardan standart ve tekrar kullanılabilir özellik üretimiydi [3,4]. 2021–2022 çalışmaları bu yaklaşımı yapılandırılabilir kohort, kalite kontrol, imputasyon ve farklı klinik görevlere genişletti [18–20]. 2023–2026 döneminde odak, tek bir veri kümesine özgü feature matrix üretiminden ortak veri katmanlarına kaydı: FEMR/EHRSHOT OMOP ve MIMIC kaynaklarını ortak arayüzde işlerken, FHIR tabanlı hatlar ve Cumulus standart API'ler üzerinden federatif analizi; MEDS ise model-kaynak bağımsız event şemasını hedefledi [22–24,29]. Bu çizgiye göre güncel bir pipeline değerlendirmesi yalnız süreyi değil semantik kapsam, kohort tekrar üretilebilirliği, provenance ve veri sızıntısı kontrolünü de açıkça raporlamalıdır.")

    add_heading(doc, "2.3. Paralel MIMIC çıkarımı ve Spark yürütme maliyeti", 2)
    add_body(doc, "Quick-MIMIC, MIMIC çıkarımında paralelliği işlevsel karşılaştırmayla birleştiren en doğrudan çalışmadır. Çok süreçli stratejide iki süreç MIMIC-III süresini yaklaşık %23, dört süreç ek %12 azaltmış; dört süreçte toplam hızlanma 1,50× olmasına karşın paralel verimlilik 0,37'ye düşmüştür. Yazarlar bu azalan getiriyi eşzamanlı I/O ile ilişkilendirmiştir [27]. MPI stratejisinde dört çekirdekte 3,53× hızlanma ve 0,88 verimlilik raporlanmış; 32 çekirdekte hızlanma artsa da verimlilik 0,26'ya gerilemiştir. Bu sonuçlar fiziksel topoloji, veri erişim yolu ve paralellik mekanizmasının birlikte raporlanması gerektiğini gösterir.")
    add_body(doc, "Spark'ın `local[K]` ve standalone modları aynı soyut DataFrame planını farklı süreç topolojilerinde yürütür. Aynı hostta standalone worker eklemek; sürücü–executor iletişimi, görev dağıtımı ve ayrı JVM maliyetlerini artırırken fiziksel disk veya bellek bant genişliği eklemez. Eş-konumlu uygulama literatürü, cache ve bellek veri yolu çekişmesinin performansı hem düşürebildiğini hem de değişkenleştirebildiğini göstermektedir [25,26]. Bu nedenle aynı host karşılaştırması gerçek strong scaling değil, sabit kaynak bütçesinde yürütme topolojisi deneyidir.")

    add_heading(doc, "2.4. Literatür sentezi ve araştırma boşluğu", 2)
    add_table(doc, ["Çalışma", "Veri/temsil", "Sistem değerlendirmesi", "Açık kalan nokta"], [
        ["MIMIC-Extract (2020)", "MIMIC-III, saatlik ICU", "Topoloji benchmarkı yok", "Spark/local–standalone etkisi"],
        ["FIDDLE (2020)", "MIMIC-III + eICU", "Ön işleme işlevselliği", "Fiziksel kaynak ve süre ayrıştırması"],
        ["MIMIC-IV Pipeline (2022)", "Görev/kohort yapılandırmalı", "Model uçtan uca", "ETL topolojisi"],
        ["EventStreamGPT (2023)", "Sürekli-zaman olay akışı", "Foundation-model veri hattı", "Kaynak-kısıtlı Spark benchmarkı"],
        ["EHRSHOT/FEMR (2023)", "OMOP/MIMIC uyumlu benchmark", "Hasta düzeyi split", "ETL yürütme topolojisi"],
        ["FHIR pipeline (2024)", "Bildirimsel FHIR özellikleri", "Dağıtık feature repository", "Tek-host çekişme deneyi"],
        ["Quick-MIMIC (2024)", "MIMIC-III/IV multimodal", "Çok süreç + MPI", "Spark local/standalone ve iş yükü granülerliği"],
        ["MEDS (2024–2026)", "Minimal event-centric standart", "Taşınabilir ekosistem", "MIMIC-III Spark topoloji ölçümü"],
        ["Bu çalışma", "Compact + vital pencere", "5 yapılandırma × n=5 + host telemetrisi", "Tek host ve sınırlı klinik modalite"],
    ], widths=[3.2, 4.4, 4.4, 4.1])
    add_caption(doc, "Tablo 1. 2020–2026 ilgili çalışmaların araştırma boşluğu açısından özeti. Süreler farklı veri ve donanım tanımları nedeniyle doğrudan karşılaştırılmamıştır.")
    add_body(doc, "Sentez, klinik veri hazırlama ile sistem performansı literatürünün çoğunlukla ayrı ilerlediğini göstermektedir. Quick-MIMIC iki alanı paralel MIMIC çıkarımında birleştirse de aynı fiziksel kaynak bütçesinde Spark yürütme modelini ve feature granülerliğini birlikte sınamaz. Bu çalışmanın araştırma boşluğu bu kesişimdedir. Çıkarım yalnız incelenen donanım ve konfigürasyonlara yöneliktir; yeni bir genel EHR standardı veya çok-hostlu ölçeklenme sonucu değildir.")

    add_heading(doc, "3. Materyal ve Yöntem", 1)
    add_body(doc, "Çalışma, sabit bir fiziksel kaynak bütçesinde kontrollü ve tekrarlı bir sistem benchmarkı olarak tasarlanmıştır. Bağımsız değişkenler yürütme yapılandırması ve özellik granülerliği; birincil bağımlı değişken uçtan uca ETL süresi, ikincil değişkenler aşama süreleri ile host CPU/bellek kullanımıdır. Veri sürümü, Parquet girdileri, ETL mantığı, Spark ayarları ve çıktı biçimleri senaryolar arasında sabit tutulmuştur. Çalışma klinik model doğruluğunu değil, aynı mantıksal dönüşümün yürütme davranışını değerlendirir.")
    add_heading(doc, "3.1. Veri kaynağı ve kapsam", 2)
    add_body(doc, "Veri kaynağı PhysioNet üzerinden erişilen MIMIC-III Clinical Database v1.4'tür [1,2]. Deney girdileri önceden hazırlanmış Parquet dosyalarıdır; ölçüm sırasında CSV→Parquet dönüşümü yapılmamıştır. Zaman kayıtlarında CHARTEVENTS için charttime, yoğun bakım kalışlarında intime/outtime alanları kullanılmıştır. Her hastane yatışı için referans başlangıç zamanı en erken ICU intime, referans bitiş zamanı en geç ICU outtime olarak tanımlanmıştır. Bu tekilleştirme, bir hadm_id altında birden fazla ICU kalışı olduğunda join sırasında satır patlamasını önlemiştir.")
    add_table(doc, ["Bileşen", "Deneyde kullanılan değer"], [
        ["Veri sürümü", "MIMIC-III v1.4; kontrollü erişim, kimliksizleştirilmiş veri"],
        ["ETL'de sayılan olay satırları", "CHARTEVENTS: 330.712.483; hadm_id boş olmayan LABEVENTS: 22.245.034"],
        ["Depolama biçimi", "Parquet (tüm timing JSON kayıtlarında doğrulandı)"],
        ["Compact çıktı", "58.976 hadm_id satırı"],
        ["Timeseries çıktı", "1.180.395 hadm_id × window_idx satırı"],
        ["Zaman penceresi", "Yalnız vital CHARTEVENTS için 6 saat; floor((charttime − intime_ref)/6 saat)"],
        ["Zaman zarfı filtresi", "min ICU intime ≤ charttime ≤ max ICU outtime; varsayılan açık"],
        ["Kırpma sonucu", "28.794.711 aday ölçümden 28.652.828 tutuldu; 141.883 (%0,493) elendi"],
    ], widths=[4.2, 12.0])
    add_caption(doc, "Tablo 2. Veri kapsamı ve iki özellik ürününün tanımı.")

    add_heading(doc, "3.2. ETL iş yükleri", 2)
    add_body(doc, "Compact rejim, vital bulgular ve laboratuvar ölçümlerini hadm_id düzeyinde özetleyerek kabul başına tek satır üretir. Vital özellikler klinik aralık filtresinden sonra ortalama ve sayı; laboratuvarlar ortalama, standart sapma, sayı ve anormallik göstergeleriyle temsil edilir. Timeseries rejiminde yalnızca seçilmiş CHARTEVENTS vital bulguları yoğun bakım referans başlangıcına göre altı saatlik pencerelere atanır ve hadm_id ile window_idx üzerinde ortalama, minimum, maksimum, standart sapma ve sayı hesaplanır. Bir pencere satırı ancak filtreleri geçen en az bir vital ölçüm içeriyorsa oluşur; düzenli ve boş pencereler üretilmez. Laboratuvar ve ICU özellikleri hadm_id düzeyinde özetlenir ve aynı kabulün bütün pencere satırlarına tekrar eklenir. Etiket ve demografik değişkenler de kabul anahtarıyla eklenir. Bu ürün bu nedenle tam çok-değişkenli klinik zaman serisi değil, pencere-seviyeli vital özetleri ile kabul-seviyeli bağlam özelliklerinin bir birleşimidir. Sayısal boşluklar ETL aşamasında keyfî olarak sıfırla doldurulmaz.")
    add_body(doc, "Zaman zarfı filtresi yalnızca timeseries rejiminde uygulanmıştır. Bir hadm_id için en erken ICU intime ile en geç ICU outtime arasında kalan vital ölçümler tutulmuştur. MIMIC veri modelinde hospital admission (hadm_id) ile ICU stay (icustay_id) farklı seviyelerdir ve aynı kabul birden fazla ICU kalışı içerebilir [1,15]. Bu tasarım join satır patlamasını önler; ancak ICU kalışları arasında ICU dışında geçirilen zamanı da kapsayabilir. Dolayısıyla filtre tekil ICU kalış üyeliğini garanti etmez ve klinik kohort üretimi için icustay_id düzeyinde yeniden tasarlanmalıdır. Raporlanan 28.794.711 aday ve 141.883 elenen satır, vital itemid, sayısal değer, klinik aralık, charttime ve stay_ref filtrelerinden sonraki tanısal sayımlardır; ham CHARTEVENTS toplamının doğrudan parçalanması değildir.")
    add_figure(doc, "sekil_1_mimari.png", "Şekil 1. ETL veri akışı ve aynı fiziksel host üzerinde karşılaştırılan yürütme senaryoları.")

    add_heading(doc, "3.3. Donanım, yazılım ve yürütme topolojisi", 2)
    add_table(doc, ["Katman", "Sürüm / yapılandırma"], [
        ["Host işletim sistemi", "Windows 11 Pro, build 26100"],
        ["İşlemci", "Intel Core i5-1335U; 10 fiziksel, 12 mantıksal çekirdek"],
        ["Bellek", "15,529 GB host RAM"],
        ["Depolama", "WD PC SN740 512 GB NVMe SSD"],
        ["Apache Spark", "3.5.1; PySpark / Spark SQL"],
        ["Docker / Compose", "Docker 29.4.0; Compose 5.1.2"],
        ["Spark ayarları", "driver 2g; executor 2g; shuffle partitions 8; AQE açık; output partitions 1"],
        ["Worker ilanı", "Her worker --cores 12 ve --memory 6g; Docker cgroup CPU/RAM kotası yok"],
        ["Paylaşılan kaynak", "Tüm senaryolar aynı CPU, RAM, SSD ve host ağ yığınını kullanır"],
    ], widths=[4.2, 12.0])
    add_caption(doc, "Tablo 3. Deney ortamı. Worker değerleri fiziksel tahsis değil, Spark'a ilan edilen üst sınırlardır.")
    add_body(doc, "local[2], local[4] ve local[8], Spark dokümantasyonundaki tanımıyla aynı süreçte sırasıyla 2, 4 ve 8 worker iş parçacığı kullanır [5]. Standalone senaryolarda master'a sırasıyla bir ve iki worker kapsayıcısı bağlanmıştır [6]. Her worker'ın 12 çekirdek ilan etmesi, 12 mantıksal çekirdekli host üzerinde iki worker senaryosunda aşırı abonelik potansiyeli yaratır. Bu nedenle '2 düğüm' ve '3 düğüm' etiketleri fiziksel sunucu sayısını değil, aynı hosttaki Spark süreç topolojisini ifade eder.")

    add_heading(doc, "3.4. Deney protokolü ve ölçümler", 2)
    add_body(doc, "Beş senaryo her iki iş yükünde altışar kez yürütülmüştür. Her senaryonun r01 koşusu ön ısıtma olarak dışlanmış, r02–r06 ölçümleri analiz edilmiştir (senaryo başına n=5; toplam 50 ölçüm koşusu). Extract süresi beş seçilmiş DataFrame'in oluşturulmasını, DISK_ONLY kalıcılığa alınmasını ve satır sayımıyla materyalize edilmesini kapsar; yalnız dosya okuma süresi değildir. Transform süresi kalıcı feature DataFrame'in kurulması ve count() ile materyalize edilmesini içerir. Load süresi aynı kalıcı sonuçtan hem Parquet hem başlıklı CSV yazılmasını ve `coalesce(1)` uygulanmasını kapsar. total_seconds bu üç uygulama içi aşamanın toplamıdır; worker başlatma/durdurma ve deney sürücüsünün Docker kontrol süresi bu değere dahil değildir. Girdi biçimi her koşunun timing JSON kaydında parquet olarak denetlenmiştir. Kaynak kullanımı psutil ile host düzeyinde yaklaşık bir saniye aralıkla örneklenmiştir; ölçümler kapsayıcı başına değildir.")
    add_body(doc, "Çıktı eşdeğerliği satır sayısı ve şema düzeyinde doğrulanmıştır: aynı iş yükündeki bütün yürütme senaryoları aynı feature kolonlarını ve aynı satır sayısını üretmiştir. Sayısal agregasyonlarda kayan nokta toplama sırasına bağlı son-bit farkları için hücre-düzeyi kriptografik eşitlik testi yapılmamıştır. Bu nedenle doğrulama mantıksal çıktı kapsamını destekler, bit düzeyi özdeşlik iddiası taşımaz.")
    add_body(doc, "Senaryo sırası deney sürücüsünde sabittir ve işletim sistemi sayfa önbelleği temizlenmemiştir. İlk koşunun dışlanması süreç/JIT ve ilk okuma etkisini azaltmakla birlikte bütün zamansal ve önbellek etkilerini ortadan kaldırmaz. Bu durum sınırlılıklar bölümünde ele alınmıştır.")

    add_heading(doc, "3.5. İstatistiksel analiz", 2)
    add_body(doc, "Her senaryo için ortalama±örnek standart sapma ve medyan [Q1–Q3] raporlandı. Varyans homojenliği medyan-merkezli Brown–Forsythe testiyle değerlendirildi. Küçük ve dengesiz varyanslara duyarlı klasik ANOVA yerine Welch tek yönlü ANOVA birincil omnibus test olarak kullanıldı; klasik ANOVA ve dağılım varsayımına daha az duyarlı Kruskal–Wallis testi duyarlılık analizleri olarak raporlandı. Önceden belirlenen local[8] referansına karşı dört Welch t-testi uygulandı; çoklu karşılaştırma için Holm düzeltmesi ve standartlaştırılmış etki büyüklüğü olarak küçük örneklem düzeltmeli Hedges g hesaplandı. İki yönlü α=0,05 kullanıldı. n=5 nedeniyle normallik testleri karar ölçütü olarak kullanılmadı ve p-değerleri betimsel dağılımlar ile birlikte yorumlandı. Ham timing CSV'leri tek doğruluk kaynağı kabul edilerek bütün özetler yeniden üretildi.")

    add_heading(doc, "4. Bulgular", 1)
    add_heading(doc, "4.1. Compact iş yükü", 2)
    add_table(doc, ["Senaryo", "Çıkarma (s)", "Dönüşüm (s)", "Yazma (s)", "Toplam (s)", "Medyan [Q1–Q3]"], runtime_rows(summary, "Compact"), widths=[2.2, 2.5, 2.5, 2.2, 2.5, 3.3])
    add_caption(doc, "Tablo 4. Compact iş yükü süreleri, ortalama±SS (n=5); r01 dışlanmıştır.")
    add_body(doc, "Compact iş yükünde en düşük ortalama toplam süre local[8] için 204,8±2,9 s'dir. local[2]'ye göre ortalama süre farkı −62,5 s (%23,4), bir worker'lı standalone senaryoya göre −22,0 s (%9,7) olmuştur. Üç süreçli topoloji 320,7±30,5 s ile en yavaş ve local[8]'e göre %56,6 daha uzun sürmüştür. Sürenin ana bileşeni bütün senaryolarda Parquet tablolarının çıkarılmasıdır.")

    add_heading(doc, "4.2. Timeseries iş yükü", 2)
    add_table(doc, ["Senaryo", "Çıkarma (s)", "Dönüşüm (s)", "Yazma (s)", "Toplam (s)", "Medyan [Q1–Q3]"], runtime_rows(summary, "Timeseries"), widths=[2.2, 2.5, 2.5, 2.2, 2.5, 3.3])
    add_caption(doc, "Tablo 5. Altı saatlik timeseries iş yükü süreleri, ortalama±SS (n=5); r01 dışlanmıştır.")
    add_body(doc, "Timeseries çıktısının satır sayısı compact matristen yaklaşık 20 kat büyüktür. En düşük ortalama süre yine local[8]'de 263,9±30,1 s'dir; medyanı 251,8 s'dir. local[4] ile tek worker'lı standalone topolojinin ortalamaları neredeyse aynıdır (291,5 ve 291,7 s). Üç düğüm senaryosu 476,8±133,0 s ile hem en yavaş hem en değişkendir; tek bir yüksek süreli koşu dağılımı genişletmiştir. Bu nedenle ortalama yanında medyan ve ham noktalar birlikte değerlendirilmiştir.")
    add_table(doc, ["Çalışma / rejim", "Paralellik karşılaştırması", "Hızlanma", "Basit verimlilik"], [
        ["Bu çalışma — compact", "local[2] → local[8]", "1,30×", "0,33"],
        ["Bu çalışma — timeseries", "local[2] → local[8]", "1,16×", "0,29"],
        ["Quick-MIMIC — MIMIC-III", "1 → 4 süreç", "1,50×", "0,37"],
        ["Quick-MIMIC — MPI", "1 → 4 çekirdek", "3,53×", "0,88"],
    ], widths=[4.0, 4.2, 3.0, 3.4])
    add_caption(doc, "Tablo 6. Paralellik arttıkça gözlenen hızlanma ve basit paralel verimlilik. Quick-MIMIC satırları farklı donanım, yazılım ve çıktı kapsamına dayandığından doğrudan performans sıralaması değildir [27].")
    add_figure(doc, "sekil_2_calisma_suresi.png", "Şekil 2. İş yükü ve yürütme senaryosuna göre ölçüm koşullarının toplam ETL süre dağılımı. Noktalar r02–r06'yı, elmaslar ortalamayı gösterir.")
    add_figure(doc, "sekil_3_asama_kirilimi.png", "Şekil 3. Extract, transform ve load aşamalarının toplam süreye ortalama katkısı; hata çubukları toplam sürenin standart sapmasıdır (n=5).")

    add_heading(doc, "4.3. İstatistiksel karşılaştırmalar", 2)
    om_rows = []
    for _, r in omnibus.iterrows():
        om_rows.append([
            r.workload,
            f"F={r.brown_forsythe_F:.2f}; p={r.brown_forsythe_p:.3g}",
            f"F({r.welch_df1:.0f},{r.welch_df2:.1f})={r.welch_F:.2f}; p={r.welch_p:.3g}",
            f"F={r.anova_F:.2f}; p={r.anova_p:.3g}",
            f"H={r.kruskal_H:.2f}; p={r.kruskal_p:.3g}",
        ])
    add_table(doc, ["İş yükü", "Brown–Forsythe", "Welch ANOVA", "Klasik ANOVA", "Kruskal–Wallis"], om_rows, widths=[2.2, 3.2, 3.8, 3.2, 3.2])
    add_caption(doc, "Tablo 7. Varyans denetimi ve beş yürütme yapılandırmasının omnibus karşılaştırmaları. Birincil sonuç Welch ANOVA'dır.")
    pair_rows = []
    for _, r in pairwise.iterrows():
        comp = r.comparison.replace("dÃ¼ÄŸÃ¼m", "düğüm")
        pair_rows.append([r.workload, comp, f"{r.mean_difference_s:.1f}", f"{r.hedges_g:.2f}", f"{r.p_raw:.4f}", f"{r.p_holm:.4f}"])
    add_table(doc, ["İş yükü", "Karşılaştırma", "Fark (s)", "Hedges g", "Ham p", "Holm p"], pair_rows, widths=[2.0, 5.0, 2.1, 2.0, 2.1, 2.1])
    add_caption(doc, "Tablo 8. local[8] referanslı Welch karşılaştırmaları. Negatif fark local[8]'in daha kısa sürdüğünü gösterir.")
    add_body(doc, "Compact iş yükünde Brown–Forsythe testi varyansların eşitliğiyle uyumsuzdur (p=0,033); bu nedenle klasik ANOVA'nın çok küçük p-değeri tek başına kullanılmamıştır. Welch ANOVA hem compact (F(4,9,1)=55,18; p<0,001) hem timeseries (F(4,9,3)=4,50; p=0,027) için yapılandırma etkisini desteklemiştir. Kruskal–Wallis sonuçları aynı yöndedir. Compact iş yükünde bütün local[8] karşılaştırmaları Holm düzeltmesinden sonra anlamlıdır (düzeltilmiş p≤0,0021) ve etki büyüklükleri yüksektir. Timeseries rejiminde local[8] ile yapılan dört ikili karşılaştırmanın hiçbiri Holm düzeltmesinden sonra α=0,05 eşiğini geçmemiştir. Bu sonuç, timeseries örnekleminde local[8]'in en düşük ortalama ve medyana sahip olduğu gözlemini değiştirmez; fakat n=5 ve yüksek değişkenlik altında ayrı ayrı üstünlük iddiasını desteklemez.")

    add_heading(doc, "4.4. Host kaynak kullanımı", 2)
    cpu_rows = []
    for _, r in cpu.iterrows():
        cpu_rows.append([SCENARIO_LABEL[r.scenario], f"{r.cpu_mean:.1f} ± {r.cpu_std:.1f}", f"{r.cpu_p95_mean:.1f}", f"{r.memory_mean:.1f}", f"{r.memory_peak_max:.1f}"])
    add_table(doc, ["Senaryo", "Ortalama CPU %", "CPU p95 %", "Ort. bellek %", "Tepe bellek %"], cpu_rows, widths=[3.0, 3.3, 3.0, 3.3, 3.3])
    add_caption(doc, "Tablo 9. Timeseries ölçüm koşularının host düzeyi kaynak özeti (ortalama±SS, n=5).")
    add_body(doc, "local[8], %81,0±0,3 ortalama host CPU kullanımıyla en yüksek ve koşular arasında en kararlı kullanımı üretmiştir. İki ve üç düğümlü topolojilerde ortalama CPU sırasıyla %74,5 ve %75,9 iken ortalama bellek kullanımı %91,0 ve %94,7'ye çıkmış, en az bir koşuda %100 tepe değer görülmüştür. Bu ölçümler süreç başına tüketimi ayırmadığından nedensel bir JVM maliyeti hesabı değildir; yine de aynı fiziksel kaynaklar üzerindeki çekişmeyle uyumludur.")
    add_figure(doc, "sekil_4_cpu_kullanimi.png", "Şekil 4. ETL koşusu boyunca host CPU kullanımı. Solda her senaryonun medyan süreye en yakın temsili koşusu; sağda r02–r06 ortalama±SS. Metrikler psutil ile host düzeyinde yaklaşık 1 sn aralıkla örneklenmiştir (kapsayıcı başına değil).")

    add_heading(doc, "5. Tartışma", 1)
    add_heading(doc, "5.1. Temel bulguların yorumu", 2)
    add_body(doc, "İki iş yükünde de en düşük gözlenen ortalama süre local[8]'e aittir. Compact iş yükünde farklar kararlı ve ikili testlerde anlamlıdır. Timeseries iş yükünde yön aynı olsa da local[8] değişkenliği artmış ve Holm-düzeltilmiş ikili kanıt yetersiz kalmıştır; bu nedenle sonuç 'local[8] kesin olarak daha hızlıdır' değil, 'bu örneklemde en düşük ortalama ve medyan local[8]'de gözlenmiştir' şeklinde ifade edilmelidir. Deney tasarımı aynı host üzerindeki yapılandırmaları karşılaştırdığı için bulgu, farklı fiziksel kaynak bütçeleri veya gerçek cluster boyutları arasındaki nedensel hızlanma olarak genellenemez.")
    add_body(doc, "Dağıtık işleme, bağımsız fiziksel kaynaklar eklenerek toplam hesaplama ve I/O kapasitesi yükseldiğinde yarar sağlar. Bu deneyde worker'lar aynı 12 mantıksal çekirdeği, 15,5 GB belleği ve tek SSD'yi paylaşmıştır. Ayrıca her worker 12 çekirdek ilan etmiş, ancak Docker cgroup kotasıyla yalıtılmamıştır. İki worker'lı senaryoda executor rekabeti, süreçler arası iletişim, serileştirme ve shuffle koordinasyonu gerçek bir kaynak artışı olmadan eklenmiştir. Üç düğüm senaryosundaki yüksek bellek kullanımı ve süre varyansı bu mekanizmayla uyumludur; ancak Spark event loglarına dayalı executor-spesifik kanıt bulunmadığı için kesin neden ataması yapılmamıştır.")
    add_body(doc, "Timeseries iş yükü çıktı satır sayısını büyük ölçüde artırmış olsa da toplam sürenin en büyük bölümü extract aşamasında kalmıştır. Extract ölçümü ham taramanın yanında DISK_ONLY persist ve beş tablo için count() eylemlerini de içerdiğinden bunu doğrudan depolama bant genişliği olarak yorumlamak mümkün değildir. Dönüşümde yalnız vital CHARTEVENTS pencerelenmiş, laboratuvarlar kabul düzeyinde kalmış ve itemid/klinik aralık filtreleri agregasyon öncesinde veri hacmini azaltmıştır. Dolayısıyla nihai satır sayısı tek başına shuffle yoğunluğunu veya beklenen dağıtık kazancı tanımlamaz. Spark SQL'in Uyarlanabilir Sorgu Yürütmesi çalışma zamanı istatistikleriyle shuffle bölümlerini birleştirebilir ve çarpık join'leri ele alabilir [11]; burada AQE açık olmakla birlikte başlangıç shuffle bölümü sayısının sekiz olması, worker başına iki executor çekirdeği ve çıktının `coalesce(1)` ile hem Parquet hem CSV'ye yazılması üst ölçek sınırları yaratmış olabilir.")

    add_heading(doc, "5.2. Güncel çalışmalarla karşılaştırma", 2)
    add_body(doc, "Quick-MIMIC'in MIMIC-III çok süreçli deneyinde dört süreç, tek sürece göre 1,50× hızlanma ve 0,37 paralel verimlilik sağlamıştır [27]. Bu çalışmada local[2]'den local[8]'e compact hızlanma 1,30×, timeseries hızlanma 1,16×'tir; iş parçacığı kapasitesi dört katına çıktığı için karşılık gelen basit paralel verimlilikler yaklaşık 0,33 ve 0,29'dur. Her iki çalışma da paralellik arttıkça azalan getiri göstermektedir. Bununla birlikte Quick-MIMIC hasta-vaka çıkarım süresini, farklı donanım ve SQL/MPI mimarisiyle ölçtüğü; bu çalışma ise iki biçimde çıktı yazan uçtan uca Spark aşamalarını ölçtüğü için mutlak süreler doğrudan karşılaştırılamaz.")
    add_body(doc, "Quick-MIMIC'in en iyi MPI sonucu dört çekirdekte 3,53× hızlanma ve 0,88 verimliliktir; bu değerler mevcut local[8] sonucundan belirgin biçimde yüksektir [27]. Ayrışmanın olası nedenleri Quick-MIMIC'in hasta-temelli görev bölme stratejisi, MPI süreç yerleşimi, farklı depolama erişimi ve mevcut hattın bütün ham tabloları her koşuda sayması, Spark görev/JVM maliyeti ve tek parçalı çift çıktı yazmasıdır. Bu yorum nedensel test değil, mimari farklara dayalı açıklamadır.")
    add_body(doc, "Eş-konumlandırma literatürü, aynı sunucuda cache, bellek denetleyicisi ve I/O paylaşımının performans kaybı ve varyans üretebildiğini raporlamaktadır [25,26]. Bu çalışmada local[8] timeseries CPU kullanımı %81,0 ile en yüksekken iki-worker topolojisinde bellek ortalaması %94,7'ye ve tepe değer %100'e çıkmıştır. Bulgular bu literatürle uyumludur; fakat süreç bazlı donanım sayaçları ve Spark event logları olmadığından çekişmenin hangi alt sistemden kaynaklandığı ayrıştırılamaz.")

    add_heading(doc, "5.3. Klinik veri mühendisliği açısından sonuçlar", 2)
    add_body(doc, "Zaman zarfı filtresi yalnızca performans ayarı değil veri geçerliliği kararıdır. Uzak zaman damgalarının pencerelere dönüşmesini azaltmıştır; ancak min(intime)–max(outtime) zarfı birden fazla ICU kalışı arasındaki olası servis dönemlerini dışlamaz. Ayrıca laboratuvar özetleri tüm kabul üzerinden hesaplandığı için ürün, belirli bir erken tahmin zamanı için sızıntısız klinik özellik matrisi değildir. Bu makale üretilen özelliklerin prognostik yararını veya klinik uygulanabilirliğini test etmemektedir. MIMIC-Extract ve FIDDLE'ın vurguladığı gibi temsil, kohort, zaman ufku ve eksik veri stratejisi model hedefiyle birlikte tanımlanmalıdır [3,4].")
    add_body(doc, "Gelecekte mortalite modeli geliştirilecekse aynı hadm_id veya subject_id'ye ait bütün pencereler aynı veri bölümünde tutulmalı; tahmin anından sonra oluşan laboratuvarlar, ICU çıkış zamanı ve toplam yatış süresi gibi geleceği taşıyan değişkenler dışlanmalıdır. Değerlendirme AUROC yanında AUPRC, kalibrasyon eğrisi/intercept/slope, Brier skoru, karar eğrisi ve klinik alt grup sonuçlarını içermeli; raporlama TRIPOD+AI ile uyumlu olmalıdır [7,8,10].")

    add_heading(doc, "5.4. Yeniden üretilebilirlik ve pratik çıkarımlar", 2)
    add_body(doc, "Docker Compose, yazılım bileşenlerini ve hizmet topolojisini kayıt altına alarak deney ortamının yeniden kurulmasını kolaylaştırır; ancak kapsayıcı kullanımı tek başına bit düzeyinde yeniden üretilebilirlik veya donanımlar arası eşdeğer performans garantisi değildir [12]. Bu çalışmada ham timing JSON/CSV dosyaları, kaynak günlükleri, analiz betiği, Spark ETL betiği ve yapılandırma dosyası sonuçların yeniden hesaplanması için gerekli minimum kanıt paketini oluşturur. Kontrollü MIMIC-III verisi ve hasta düzeyi ürünler veri kullanım sözleşmesi nedeniyle dağıtılamaz.")
    add_body(doc, "Pratik olarak, tek fiziksel makinede Spark worker kapsayıcısı sayısını artırmak kaynak artışı olarak kabul edilmemelidir. Önce local iş parçacığı sayısı, veri bölümleri ve seri yazma noktaları profillenmeli; standalone süreç topolojisi ancak süreç izolasyonu, hata toleransı veya gerçek uzak worker ihtiyacı varsa seçilmelidir. Genellenebilir performans iddiası için worker'ların ayrı fiziksel hostlarda, eşit ve doğrulanmış cgroup kaynaklarıyla ve Spark event-log telemetrisi altında çalıştırılması gerekir.")

    add_heading(doc, "6. Geçerlilik Tehditleri ve Sınırlılıklar", 1)
    add_heading(doc, "6.1. İç geçerlilik", 2)
    add_body(doc, "Senaryolar sabit sırada çalıştırılmış ve işletim sistemi disk önbelleği temizlenmemiştir. r01'in dışlanması JIT, süreç başlatma ve ilk okuma etkilerini azaltabilir; ancak sıra, sıcak önbellek, arka plan süreçleri ve termal durumun yapılandırmayla karışmasını engellemez. Worker kapsayıcılarında cgroup CPU/RAM kotası bulunmaması ve her worker'ın 12 mantıksal çekirdeğin tamamını ilan etmesi özellikle iki-worker koşulunda aşırı abonelik yaratır. Bu nedenle gözlenen farklar yalnız Spark topolojisine değil, topolojiyle birlikte oluşan kaynak rekabetine aittir.")
    add_heading(doc, "6.2. Yapı geçerliliği", 2)
    add_body(doc, "`extract_seconds` saf disk-okuma süresi değildir; DataFrame oluşturma, DISK_ONLY persist ve beş count() eylemini kapsar. `load_seconds`, `coalesce(1)` sonrasında hem Parquet hem CSV yazmayı içerir. Host psutil telemetrisi driver, executor ve diğer süreçleri ayırmaz; Spark event logu, GC, shuffle spill, görev gecikmesi ve disk I/O sayaçları sistematik toplanmamıştır. Dolayısıyla CPU/bellek örüntüleri mekanizmayla uyumlu kanıt sağlar, belirli bir Spark alt sistemine nedensel atama sağlamaz.")
    add_heading(doc, "6.3. Sonuç geçerliliği", 2)
    add_body(doc, "Senaryo başına n=5, özellikle timeseries iki-worker koşulundaki yüksek varyansı hassas tahmin etmek için sınırlıdır. Welch ANOVA heteroskedastisiteyi azaltmak için seçilmiş, Holm düzeltmesi aile-düzeyi hatayı denetlemiştir; buna rağmen küçük örneklem güven aralıklarını genişletir. Timeseries omnibus sonucu anlamlı olsa da local[8] ikili karşılaştırmalarının hiçbiri düzeltme sonrasında anlamlı değildir. Bu nedenle sıralama betimsel, üstünlük iddiası ise yalnız compact rejim için güçlüdür.")
    add_heading(doc, "6.4. Dış geçerlilik", 2)
    add_body(doc, "Deney tek bir Windows dizüstü host, tek SSD ve tek Spark sürümüyle yürütülmüştür. Sonuçlar ayrı ağ, disk ve bellek kaynakları bulunan gerçek çok-hostlu kümeler için strong-scaling kanıtı değildir. MIMIC-III tek merkezli ve tarihsel bir kaynaktır; çalışma klinik tahmin performansı, adalet, dış doğrulama veya hasta yararı iddiasında bulunmaz. `output_partitions=1` ve sekiz shuffle bölümü başka ölçeklerde farklı darboğazlar yaratabilir.")
    add_heading(doc, "6.5. Klinik temsil ve yeniden üretilebilirlik", 2)
    add_body(doc, "Timeseries ürünü tam düzenli bir grid değildir; yalnız uygun vital olayı bulunan hadm_id–window_idx çiftleri oluşur ve laboratuvar/ICU özetleri kabul düzeyinde tekrarlanır. min(intime)–max(outtime) zarfı, bir kabuldeki ICU kalışları arasındaki ICU-dışı zamanı kapsayabilir. Bu nedenle ürün, klinik tahmin için sızıntısız hazır veri kümesi olarak sunulmamaktadır. Kod, yapılandırma, anonim timing kayıtları ve analiz betikleri paylaşılabilir; kontrollü MIMIC verisi ve hasta düzeyi türevleri paylaşılamaz.")

    add_heading(doc, "7. Sonuç", 1)
    add_body(doc, "Aynı fiziksel host üzerinde yürütülen bu MIMIC-III Spark ETL deneyinde, compact ve vital bulguları altı saatlik pencerelerde özetleyen timeseries iş yüklerinin her ikisinde de en düşük gözlenen çalışma süresi local[8] yapılandırmasında elde edilmiştir. Ek worker kapsayıcıları fiziksel kaynak eklememiş; bir-worker standalone topolojisi local[8]'i geçmemiş, iki-worker topolojisi ise daha yavaş ve daha değişken olmuştur. Compact bulgusu düzeltilmiş ikili testlerle güçlüdür. Timeseries için Welch omnibus testi yapılandırma farkını desteklese de düzeltilmiş ikili belirsizlik korunmalıdır. Sonuçlar, tek-host kapsayıcı topolojisinin gerçek dağıtık strong scaling olarak yorumlanmaması, ETL aşama tanımlarının açık verilmesi ve negatif ölçeklenme sonuçlarının da raporlanması gerektiğini ortaya koymaktadır.")

    add_heading(doc, "Etik, veri erişimi ve çıkar çatışması bildirimleri", 1)
    add_body(doc, "Veri erişimi: MIMIC-III kontrollü erişimli, kimliksizleştirilmiş bir veri kaynağıdır. Çalışmanın yazarı/yazarları PhysioNet veri kullanım koşullarına uymalıdır. Ham veri veya hasta düzeyi türevler bu çalışma ile paylaşılmayacaktır.")
    add_body(doc, "Etik kurul / muafiyet: [Yazarın kurumu tarafından verilen etik kurul kararı, muafiyet numarası veya 'ikincil kimliksizleştirilmiş veri analizi' değerlendirmesi buraya eklenmelidir. Bu bilgi doğrulanmadan gönderim yapılmamalıdır.]")
    add_body(doc, "Çıkar çatışması: [Yazar beyanı eklenecek.]")
    add_body(doc, "Finansman: [Finansman yoksa 'Bu çalışma özel bir fon almamıştır' yazılacak; varsa destekçi ve proje numarası eklenecek.]")
    add_body(doc, "Yazar katkıları: [CRediT rollerine göre eklenecek.]")
    add_body(doc, "Kod ve toplulaştırılmış sonuçların erişilebilirliği: [Kalıcı depo bağlantısı/DOI eklenecek.] Paylaşım paketi hasta düzeyi veri içermez; kaynak kod, anonimleştirilmiş timing/resource günlükleri ve istatistik üretim betiklerini içerir.")

    add_heading(doc, "Kaynaklar", 1)
    refs = [
        "Johnson AEW, Pollard TJ, Shen L, et al. MIMIC-III, a freely accessible critical care database. Scientific Data. 2016;3:160035. doi:10.1038/sdata.2016.35.",
        "Johnson AEW, Pollard TJ, Mark RG. MIMIC-III Clinical Database (version 1.4). PhysioNet. 2016. doi:10.13026/C2XW26.",
        "Wang S, McDermott MBA, Chauhan G, Ghassemi M, Hughes MC, Naumann T. MIMIC-Extract: a data extraction, preprocessing, and representation pipeline for MIMIC-III. Proceedings of ACM CHIL. 2020:222–235. doi:10.1145/3368555.3384469.",
        "Tang S, Davarmanesh P, Song Y, Koutra D, Sjoding MW, Wiens J. Democratizing EHR analyses with FIDDLE: a flexible data-driven preprocessing pipeline for structured clinical data. J Am Med Inform Assoc. 2020;27(12):1921–1934. doi:10.1093/jamia/ocaa139.",
        "Apache Spark. Submitting Applications: Master URLs, local[K] and spark://. Spark 3.5 documentation. https://spark.apache.org/docs/3.5.8/submitting-applications.html (erişim: 20 Temmuz 2026).",
        "Apache Spark. Spark Standalone Mode. Spark 3.5 documentation. https://spark.apache.org/docs/3.5.6/spark-standalone.html (erişim: 20 Temmuz 2026).",
        "Harutyunyan H, Khachatrian H, Kale DC, Ver Steeg G, Galstyan A. Multitask learning and benchmarking with clinical time series data. Scientific Data. 2019;6:96. doi:10.1038/s41597-019-0103-9.",
        "Collins GS, Moons KGM, Dhiman P, et al. TRIPOD+AI statement: updated guidance for reporting clinical prediction models that use regression or machine learning methods. BMJ. 2024;385:e078378. doi:10.1136/bmj-2023-078378.",
        "Van Calster B, McLernon DJ, van Smeden M, Wynants L, Steyerberg EW. Calibration: the Achilles heel of predictive analytics. BMC Medicine. 2019;17:230. doi:10.1186/s12916-019-1466-7.",
        "Benchimol EI, Smeeth L, Guttmann A, et al. The REporting of studies Conducted using Observational Routinely-collected health Data (RECORD) statement. PLoS Med. 2015;12(10):e1001885. doi:10.1371/journal.pmed.1001885.",
        "Apache Spark. Performance Tuning: Adaptive Query Execution. Spark 3.5 documentation. https://spark.apache.org/docs/3.5.5/sql-performance-tuning.html (erişim: 20 Temmuz 2026).",
        "Boettiger C. An introduction to Docker for reproducible research. ACM SIGOPS Operating Systems Review. 2015;49(1):71–79. doi:10.1145/2723872.2723882.",
        "Zaharia M, Xin RS, Wendell P, et al. Apache Spark: a unified engine for big data processing. Communications of the ACM. 2016;59(11):56–65. doi:10.1145/2934664.",
        "Amdahl GM. Validity of the single processor approach to achieving large scale computing capabilities. AFIPS Spring Joint Computer Conference. 1967:483–485. doi:10.1145/1465482.1465560.",
        "MIT Laboratory for Computational Physiology. Querying MIMIC-III in Postgres: ICU stay definitions and transfers. https://mimic.mit.edu/docs/III/tutorials/intro-to-mimic-iii.html (erişim: 20 Temmuz 2026).",
        "Johnson AEW, Bulgarelli L, Shen L, et al. MIMIC-IV, a freely accessible electronic health record dataset. Scientific Data. 2023;10:1. doi:10.1038/s41597-022-01899-x.",
        "McDermott MBA, Wang S, Marinsek N, Ranganath R, Foschini L, Ghassemi M. Reproducibility in machine learning for health research: Still a ways to go. Science Translational Medicine. 2021;13(586):eabb1655. doi:10.1126/scitranslmed.abb1655.",
        "Jarrett D, Yoon J, van der Schaar M. Clairvoyance: A unified, end-to-end autoML pipeline for medical time series. International Conference on Learning Representations. 2021.",
        "Mandyam A, Liao KP, Choi JW, et al. COP-E-CAT: Cleaning and organization pipeline for EHR computational and analytic tasks. Proceedings of ACM CHIL. 2021. doi:10.1145/3459930.3469536.",
        "Gupta M, Gallamoza B, Cutrona N, Dhakal P, Poulain R, Beheshti R. An extensive data processing pipeline for MIMIC-IV. Proceedings of Machine Learning for Health. 2022. doi:10.48550/arXiv.2204.13841.",
        "McDermott MBA, Yarnell CJ, Laird D, et al. EventStreamGPT: A data pre-processing and modeling library for generative, pre-trained transformers over continuous-time patient event streams. NeurIPS Datasets and Benchmarks. 2023.",
        "Wornow M, Thapa R, Steinberg E, et al. EHRSHOT: An EHR benchmark for few-shot evaluation of foundation models. NeurIPS Datasets and Benchmarks. 2023.",
        "Namli T, et al. A declarative feature engineering pipeline for healthcare data using FHIR. Frontiers in Medicine. 2024;11:1393123. doi:10.3389/fmed.2024.1393123.",
        "Medical Event Data Standard Working Group. Medical Event Data Standard (MEDS): Facilitating machine learning for health. NEJM AI. 2026. doi:10.1056/AIra2501253.",
        "Zacarias FV, et al. Performance interference-aware resource management for co-located applications in cloud environments. Journal of Parallel and Distributed Computing. 2021. doi:10.1016/j.jpdc.2021.02.010.",
        "Gugnani S, Lu X. Characterizing and understanding resource usage in containerized data-intensive systems. arXiv. 2023. arXiv:2311.07818.",
        "Dou Y, et al. Quick-MIMIC: A multimodal data extraction pipeline for MIMIC-III and MIMIC-IV. Big Data Mining and Analytics. 2024;7(4):1333–1346. doi:10.26599/BDMA.2024.9020024.",
        "Yèche H, Kuznetsova R, Zimmermann M, et al. HiRID-ICU-Benchmark—A comprehensive machine learning benchmark on high-resolution ICU data. NeurIPS Datasets and Benchmarks. 2021.",
        "Mandel JC, et al. Cumulus: A federated clinical data platform based on FHIR. Journal of the American Medical Informatics Association. 2024. doi:10.1093/jamia/ocae130.",
    ]
    for i, ref in enumerate(refs, 1):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.65)
        p.paragraph_format.first_line_indent = Cm(-0.65)
        p.paragraph_format.space_after = Pt(3)
        p.add_run(f"{i}. {ref}")

    add_heading(doc, "Gönderim öncesi zorunlu yazar kontrolü", 1)
    for item in [
        "Yazar adı, kurum, e-posta ve sorumlu yazar bilgilerini doldurun.",
        "Hedef derginin yazım şablonunu, sözcük/şekil/tablo sınırlarını ve kaynak stilini uygulayın.",
        "PhysioNet yetkilendirmesi ile kurumun etik kurul/muafiyet beyanını doğrulayın.",
        "Kod deposu bağlantısını veya arşiv DOI'sini ekleyin; MIMIC-III ham/türev hasta verisini yüklemeyin.",
        "Mümkünse sabit sıra etkisini azaltmak için senaryoları randomize ederek ek doğrulama serisi ve Spark event-log telemetrisi toplayın.",
        "Klinik mortalite modeli ayrıca sunulacaksa subject_id/hadm_id düzeyinde ayrım, tahmin zamanı ve sızıntısız özellik tanımını yeniden uygulayın.",
    ]:
        add_bullet(doc, item)

    doc.core_properties.title = "Kaynak Kısıtlı Tek-Host Ortamında Klinik Spark ETL"
    doc.core_properties.subject = "İş yükü yoğunluğu ve yürütme topolojisinin deneysel karakterizasyonu"
    doc.core_properties.keywords = "MIMIC-III, Spark, ETL, timeseries, Docker"
    doc.core_properties.comments = "Kaynak taslak korunarak doğrulanmış proje artefaktlarından yeniden yazıldı."
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
