from __future__ import annotations

from pathlib import Path

import pandas as pd
from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
SRC = Path(__file__).resolve().parent / "MIMICIII_Spark_ETL_Makale_Literatur_Temelli_v8_final.docx"
DST = Path(__file__).resolve().parent / "MIMICIII_Spark_ETL_Makale_Literatur_Temelli_v9.docx"
VAL_ANALYSIS = ROOT / "outputs" / "validation" / "analysis"
FIG5 = VAL_ANALYSIS / "validation_runtime_distribution.png"

SUPER = {"0": "⁰", "1": "¹", "2": "²", "3": "³", "4": "⁴",
         "5": "⁵", "6": "⁶", "7": "⁷", "8": "⁸", "9": "⁹", "-": "⁻"}
HEADER_FILL = "1F4E79"
ZEBRA_FILL = "EAF2F8"


def tr(x: float, decimals: int = 1) -> str:
    return f"{x:.{decimals}f}".replace(".", ",")


def sci_p(p: float) -> str:
    if p <= 0:
        return "p<0,0001"
    if p >= 0.0001:
        return f"p={tr(p, 4)}"
    import math
    exp = math.floor(math.log10(abs(p)))
    mantissa = p / (10 ** exp)
    mant_str = tr(mantissa, 2)
    exp_str = "".join(SUPER[c] for c in str(exp))
    return f"p={mant_str}×10{exp_str}"


def shade_cell(cell, fill: str) -> None:
    tcPr = cell._tc.get_or_add_tcPr()
    shd = tcPr.makeelement(qn("w:shd"), {qn("w:val"): "clear", qn("w:color"): "auto", qn("w:fill"): fill})
    tcPr.append(shd)


def set_cell_text(cell, text: str, bold: bool = False, white: bool = False, size: int = 8) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.name = "Arial"
    run.font.bold = bold
    if white:
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    p.paragraph_format.space_after = Pt(0)


def build_table(document: Document, anchor_p, headers: list[str], rows: list[list[str]], col_widths=None):
    n_cols = len(headers)
    table = document.add_table(rows=1, cols=n_cols)
    table.style = document.styles["Table Grid"]
    for j, h in enumerate(headers):
        cell = table.rows[0].cells[j]
        set_cell_text(cell, h, bold=True, white=True)
        shade_cell(cell, HEADER_FILL)
    for i, row in enumerate(rows):
        tr_cells = table.add_row().cells
        for j, val in enumerate(row):
            set_cell_text(tr_cells[j], val, bold=False, white=False)
            if i % 2 == 0:
                shade_cell(tr_cells[j], ZEBRA_FILL)
    anchor_p._p.addprevious(table._tbl)
    return table


def main() -> None:
    document = Document(SRC)
    paras = document.paragraphs

    p_36_anchor = paras[62]     # heading "4. Bulgular" -> insert 3.6 right before it
    p_45_anchor = paras[89]     # heading "5. Tartışma" -> insert 4.5 right before it
    p_51_interp = paras[91]     # 5.1 body
    p_61_intval = paras[106]    # 6.1 body
    p_63_concl = paras[110]     # 6.3 body
    p_conclusion = paras[116]   # 7. Sonuç body
    p_checklist_rand = paras[159]  # checklist bullet about randomization

    summary = pd.read_csv(VAL_ANALYSIS / "validation_summary.csv")
    contrasts = pd.read_csv(VAL_ANALYSIS / "validation_paired_contrasts.csv")

    order = ["local2", "local4", "local8", "standalone1", "standalone2"]
    label = {"local2": "local[2]", "local4": "local[4]", "local8": "local[8]",
             "standalone1": "1-worker standalone", "standalone2": "2-worker standalone"}
    wl_label = {"compact": "Compact", "timeseries": "Timeseries"}

    # ---------- 3.6 Methods subsection ----------
    p_36_anchor.insert_paragraph_before("3.6. Randomize, kaynak-kotalı doğrulama serisi", style="Heading 2")
    p_36_anchor.insert_paragraph_before(
        "Birincil seride sabit senaryo sırası, doğrulanmamış kapsayıcı kaynak sınırları ve Spark "
        "event-log eksikliğini gidermek için ikinci, bağımsız bir doğrulama serisi önceden kayıtlı "
        "bir protokolle (randomize tam blok tasarımı, seed=20260720) yürütülmüştür. Beş yürütme "
        "senaryosu (local[2], local[4], local[8], bir-worker standalone, iki-worker standalone) ve "
        "iki iş yükü (compact, timeseries) 10 hücrelik bir tam blok oluşturmuş; blok içi hücre sırası "
        "her tekrarda yeniden karıştırılmıştır. Bir ısınma bloğunu (n=1, analize dahil edilmemiştir) "
        "12 ölçüm bloğu izlemiş, toplam 130 koşunun tamamı hatasız tamamlanmıştır (hücre başına n=12, "
        "sıfır başarısız Spark görevi). Her hücrede Docker kapsayıcı kaynakları docker update ile "
        "toplam 8 mantıksal CPU ve 8 GiB RAM'e sabitlenmiş ve her koşu öncesi docker inspect ile "
        "doğrulanmıştır; standalone yapılandırmalarında worker(lar) toplam 6 çekirdek ilan etmiş, bu "
        "da Bölüm 3.3'teki birincil seriyle aynı kaynak bütçesi mantığını korumuştur. spark.eventLog."
        "enabled=true ile üretilen Spark olay günlükleri her koşu için ayrıştırılmış; toplam executor "
        "CPU süresi, JVM GC süresi, shuffle okuma/yazma ve disk sızıntısı (spill) çıkarılmıştır. "
        "Birincil karşılaştırma, her iş yükünde local[8]'e karşı diğer dört senaryonun blok içi "
        "eşleştirilmiş örneklem t testidir (bloklama, aynı tekrar içinde aynı koşullarda ölçülen "
        "çiftleri eşler); çoklu karşılaştırma için Holm düzeltmesi iş yükü içinde uygulanmıştır.",
        style="Normal",
    )

    # ---------- 4.5 Results subsection ----------
    p_45_anchor.insert_paragraph_before("4.5. Doğrulama serisi bulguları", style="Heading 2")
    p_45_anchor.insert_paragraph_before(
        "Doğrulama serisinde de en düşük ortalama süre her iki iş yükünde local[8]'e aittir: compact "
        "için 212,5±11,1 s, timeseries için 234,8±7,3 s (n=12; Tablo 10). Bu değerler "
        "birincil seriye (Tablo 4–5) yakın ancak birebir aynı değildir; fark randomize sıra, "
        "cgroup ile zorunlu kaynak eşitleme ve artırılmış örneklem büyüklüğünden (n=5→n=12) "
        "kaynaklanmaktadır. Sıfır görev başarısız olmuş, tüm 130 koşuda çıktı satır sayısı ve girdi "
        "biçimi doğrulanmıştır.",
        style="Normal",
    )
    rows10 = []
    for _, r in summary.iterrows():
        rows10.append([
            wl_label[r["workload"]], label[r["scenario"]], str(int(r["n"])),
            f"{tr(r['total_mean'],1)}±{tr(r['total_std'],1)}",
            tr(r["total_median"], 1),
            tr(r["extract_mean"], 1), tr(r["transform_mean"], 1), tr(r["load_mean"], 1),
            tr(r["executor_cpu_s"], 1), tr(r["gc_s"], 1),
        ])
    build_table(
        document, p_45_anchor,
        ["İş yükü", "Senaryo", "n", "Toplam ort.±SS (s)", "Medyan (s)",
         "Extract ort. (s)", "Transform ort. (s)", "Load ort. (s)",
         "Executor CPU (s)", "GC (s)"],
        rows10,
    )
    p_45_anchor.insert_paragraph_before("", style="Normal")
    p_45_anchor.insert_paragraph_before(
        "Tablo 10. Doğrulama serisi (randomize, kaynak-kotalı, n=12/hücre) süre ve Spark event-log "
        "özet istatistikleri. Executor CPU ve GC süreleri koşu başına ortalamadır.",
        style="Caption",
    )
    p_45_anchor.insert_paragraph_before(
        "Blok içi eşleştirilmiş t testleri ve Holm düzeltmesi (Tablo 11), local[8]'e karşı dört "
        "alternatif senaryonun tamamında — hem compact hem timeseries iş yükünde — anlamlı "
        "fark göstermiştir (tüm p_Holm<0,0001). Bu, birincil seride yalnızca compact iş yükünde "
        "gözlenen anlamlılığın timeseries için doğrulanamamasını (Bölüm 4.3, Tablo 8) gidermektedir: "
        "artırılmış örneklem büyüklüğü ve sabit sıra etkisinden arındırılmış tasarım, timeseries "
        "local[8] üstünlüğünü de istatistiksel olarak doğrulamıştır.",
        style="Normal",
    )
    rows11 = []
    for _, r in contrasts.iterrows():
        scen = r["comparison"].split(" - ")[0]
        rows11.append([
            wl_label[r["workload"]], f"{label[scen]} − local[8]", str(int(r["n_blocks"])),
            f"{tr(r['mean_difference_s'],1)} [{tr(r['difference_ci_low'],1)}–{tr(r['difference_ci_high'],1)}]",
            f"{tr(r['geometric_runtime_ratio'],2)} [{tr(r['ratio_ci_low'],2)}–{tr(r['ratio_ci_high'],2)}]",
            sci_p(r["p_holm_within_workload"]),
        ])
    build_table(
        document, p_45_anchor,
        ["İş yükü", "Karşılaştırma", "n (blok)", "Ort. fark [%95 GA] (s)",
         "Geometrik oran [%95 GA]", "p (Holm)"],
        rows11,
    )
    p_45_anchor.insert_paragraph_before("", style="Normal")
    p_45_anchor.insert_paragraph_before(
        "Tablo 11. Doğrulama serisinde local[8] referanslı, blok içi eşleştirilmiş örneklem t testi "
        "karşılaştırmaları; p değerleri iş yükü içinde Holm düzeltmelidir. Pozitif ortalama fark ve "
        "1'in üzerindeki geometrik oran, ilgili senaryonun local[8]'den daha yavaş olduğunu gösterir.",
        style="Caption",
    )
    p_45_anchor.insert_paragraph_before(
        "Spark event-log kanıtı, gözlenen yavaşlamanın disk sızıntısı veya görev hatasından "
        "kaynaklanmadığını doğrudan göstermektedir: toplam shuffle okuma/yazma ve disk sızıntısı tüm "
        "hücrelerde ihmal edilebilir düzeyde kalmış (<0,2 GB) ve sıfır görev başarısız olmuştur. "
        "Bulgu, Bölüm 3.3'te belirtildiği gibi standalone yapılandırmalarının toplam 8 çekirdeklik "
        "bütçenin yalnızca 6'sını worker'a ayırmasıyla (driver/master için 2 çekirdek ayrılması) ve "
        "ek süreç/JVM koordinasyon gideriyle tutarlıdır.",
        style="Normal",
    )
    if FIG5.exists():
        fig_p = p_45_anchor.insert_paragraph_before("", style="Normal")
        fig_p.add_run().add_picture(str(FIG5), width=Pt(430))
    p_45_anchor.insert_paragraph_before(
        "Şekil 5. Randomize, kaynak-kotalı doğrulama deneyinde iş yükü ve senaryoya göre toplam ETL "
        "süresi dağılımı (n=12); noktalar ham koşuları gösterir.",
        style="Caption",
    )

    # ---------- targeted appends to existing sections ----------
    p_51_interp.add_run(
        " Bölüm 4.5'te raporlanan bağımsız, randomize ve kaynak-kotalı doğrulama serisi bu "
        "belirsizliği gidermiştir: artırılmış örneklem (n=12) ve sabit sıra etkisinin ortadan "
        "kaldırılmasıyla local[8] üstünlüğü timeseries iş yükünde de Holm düzeltmeli ikili "
        "testlerde anlamlı bulunmuştur."
    )
    p_61_intval.add_run(
        " Bu iki tehdit (sabit sıra ve doğrulanmamış kapsayıcı kotası), Bölüm 3.6'da tanımlanan "
        "ikinci, randomize ve cgroup ile 8 CPU/8 GiB'a sabitlenmiş bağımsız bir doğrulama serisiyle "
        "ayrıca sınanmış ve birincil seriyle aynı yönde sonuç elde edilmiştir (Bölüm 4.5)."
    )
    p_63_concl.add_run(
        " Bağımsız doğrulama serisinde (n=12, randomize sıra) bu sınırlama ortadan kalkmış ve "
        "timeseries için de local[8] karşısındaki dört ikili karşılaştırmanın tamamı Holm "
        "düzeltmesinden sonra anlamlı bulunmuştur (Bölüm 4.5, Tablo 11); dolayısıyla üstünlük "
        "iddiası artık yalnızca compact rejimle sınırlı değildir."
    )
    p_conclusion.add_run(
        " Bu sonuç, sabit sıra etkisini ortadan kaldıran, kaynakları cgroup ile eşitleyen ve Spark "
        "event-log telemetrisiyle desteklenen bağımsız bir doğrulama serisinde (n=12/hücre, 130 "
        "koşu, sıfır hata) her iki iş yükü için de istatistiksel olarak doğrulanmıştır."
    )

    p_checklist_rand.runs[0].text = (
        "Tamamlandı: Sabit sıra etkisini ortadan kaldıran randomize, kaynak-kotalı bağımsız "
        "doğrulama serisi (n=12/hücre, 130 koşu, Spark event-log telemetrisiyle) yürütüldü ve "
        "Bölüm 3.6/4.5'te raporlandı."
    )
    for extra_run in p_checklist_rand.runs[1:]:
        extra_run.text = ""

    document.save(DST)
    print(f"Saved: {DST}")


if __name__ == "__main__":
    main()
