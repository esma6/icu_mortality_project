"""Replace the stale Fig. 3 (pre-recalibration Brier scores) embedded in
the manuscript with image3_new.png (regenerated from the current,
patient-grouped-calibration cv_metrics_summary.csv by regenerate_fig3.py),
and note in the caption that Brier scores are post-recalibration.
"""
import sys
sys.stdout.reconfigure(encoding="utf-8")
from docx import Document
from docx.oxml.ns import qn

DOC = "IEEE_JBHI_Submission_Manuscript.docx"
NEW_IMAGE = "image3_new.png"

d = Document(DOC)
paras = d.paragraphs

# Find the image paragraph immediately preceding the Fig. 3 caption.
fig3_caption = None
for i, p in enumerate(paras):
    if p.text.startswith("Fig. 3. Five-fold patient-grouped cross-validation results"):
        fig3_caption = i
        break
assert fig3_caption is not None, "Fig. 3 caption not found"

image_para = paras[fig3_caption - 1]
blips = image_para._p.findall(".//" + qn("a:blip"))
assert len(blips) == 1, f"expected exactly one image in the preceding paragraph, found {len(blips)}"
blip = blips[0]
rid = blip.get(qn("r:embed"))

image_part = d.part.related_parts[rid]
with open(NEW_IMAGE, "rb") as f:
    image_part._blob = f.read()

# Update caption to clarify Brier scores are post-recalibration.
cap = paras[fig3_caption]
old_cap = cap.text
assert old_cap == (
    "Fig. 3. Five-fold patient-grouped cross-validation results. Error bars "
    "denote standard deviations across folds; lower Brier scores are better."
), old_cap
new_cap = (
    "Fig. 3. Five-fold patient-grouped cross-validation results, after "
    "within-training-fold, patient-grouped recalibration (Section III.D). "
    "Error bars denote standard deviations across folds; lower Brier scores "
    "are better."
)
for run in list(cap.runs)[1:]:
    run.text = ""
if cap.runs:
    cap.runs[0].text = new_cap
else:
    cap.add_run(new_cap)

d.save(DOC)
print("Replaced Fig. 3 image and updated caption.")
print(new_cap)
