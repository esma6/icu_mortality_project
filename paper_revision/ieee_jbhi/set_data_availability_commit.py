"""Update the Data and Code Availability commit hash. Run once, after
committing the code/data fix that the manuscript's numbers depend on, with
that commit's hash as the sole argument.

Usage: python set_data_availability_commit.py <new_hash>
"""
import re
import sys

sys.stdout.reconfigure(encoding="utf-8")
from docx import Document

DOC = "IEEE_JBHI_Submission_Manuscript.docx"
HASH_RE = re.compile(r"commit [0-9a-f]{40}")


def main() -> None:
    if len(sys.argv) != 2 or not re.fullmatch(r"[0-9a-f]{40}", sys.argv[1]):
        raise SystemExit("usage: python set_data_availability_commit.py <40-char-hex-commit-hash>")
    new_hash = sys.argv[1]

    d = Document(DOC)
    target = None
    for p in d.paragraphs:
        if HASH_RE.search(p.text):
            target = p
            break
    if target is None:
        raise SystemExit("No 'commit <hash>' paragraph found")

    new_text = HASH_RE.sub(f"commit {new_hash}", target.text)
    for run in list(target.runs)[1:]:
        run.text = ""
    if target.runs:
        target.runs[0].text = new_text
    else:
        target.add_run(new_text)
    d.save(DOC)
    print(new_text)


if __name__ == "__main__":
    main()
