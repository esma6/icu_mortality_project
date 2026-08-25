"""Rebuild the IEEE JBHI submission manuscript against the corrected pipeline
output (landmark cohort fix, lab feature fix, calibration fix, local[6]
equal-capacity control -- see the "Apply pre-review remediation" commit).

The original draft (as received) cited git commit 55c8926, which predates all
of these fixes: its headline numbers (31,252 admissions, 12.58% mortality,
uncalibrated probabilities with a large negative intercept, no equal-capacity
control for the local[8]-vs-standalone core-count confound) are the same
numbers the BMC pre-review report flagged as requiring major revision.

Every number substituted below was read from the current (fixed, committed)
outputs/tables_ml_leakfree/*.csv and outputs/validation/analysis/*.csv files,
not retyped from memory. Run from this directory against a fresh copy of the
original IEEE_JBHI_Submission_Manuscript.docx.
"""
import sys
sys.stdout.reconfigure(encoding="utf-8")
from docx import Document

DOC = "IEEE_JBHI_Submission_Manuscript.docx"


def replace_paragraph_text(paragraph, new_text):
    for run in list(paragraph.runs)[1:]:
        run.text = ""
    if paragraph.runs:
        paragraph.runs[0].text = new_text
    else:
        paragraph.add_run(new_text)


def fix_substring(paragraph, old, new, label):
    text = paragraph.text
    if old not in text:
        raise ValueError(f"[{label}] not found: {old!r}\nIn: {text[:200]!r}")
    replace_paragraph_text(paragraph, text.replace(old, new))


def main():
    d = Document(DOC)
    paras = d.paragraphs  # captured once; stable references survive later insertions
    tables = d.tables

    # ---------------------------------------------------------------
    # Abstract
    # ---------------------------------------------------------------
    assert paras[3].text.startswith("Abstract—Reproducible clinical modeling"), paras[3].text[:80]
    replace_paragraph_text(
        paras[3],
        "Abstract—Reproducible clinical modeling depends on both clinically "
        "valid feature definitions and transparent computational execution. We "
        "developed an Apache Spark pipeline for MIMIC-III that separates two "
        "purposes often conflated in electronic-health-record processing: "
        "systems benchmarking and prediction-time feature construction. Two "
        "whole-admission outputs were used only to compare five Spark "
        "configurations on one physical host. A separate feature set used "
        "measurements from the first 48 h of intensive-care admission, excluded "
        "discharge-adjacent variables, and was evaluated for in-hospital "
        "mortality with patient-grouped splits. In a randomized, resource-capped "
        "experiment (12 paired blocks; 8 CPUs and 8 GiB), local[8] had the "
        "shortest mean runtime for the admission-level (212.5 ± 11.1 s) and "
        "six-hour-window (234.8 ± 7.3 s) workloads. All eight prespecified "
        "contrasts with local[8] remained significant after Holm adjustment; "
        "exact paired sensitivity tests gave p = 0.000488. Because standalone "
        "configurations declared six executor cores while local[8] exposed "
        "eight task threads, a supplementary equal-capacity comparison "
        "(local[6] vs. a two-worker standalone arrangement) was run in the same "
        "session; local[6] remained faster for both workloads (six-hour window: "
        "p < 0.0001; admission-level: p = 0.053). The leakage-aware cohort, "
        "redefined around each admission's first ICU stay rather than a "
        "multi-stay envelope, contained 29,886 admissions from 25,271 patients "
        "with 12.4% mortality. Gradient boosting achieved an area under the "
        "receiver-operating-characteristic curve of 0.815 ± 0.005 and an area "
        "under the precision-recall curve of 0.402 ± 0.025 in five-fold "
        "patient-grouped cross-validation; within-training-fold recalibration "
        "of the class-weighted probabilities corrected a large negative "
        "calibration intercept to 0.092 ± 0.054 and yielded a Brier skill "
        "score of 0.169 relative to a constant-prevalence baseline. These "
        "results show that co-located Spark workers do not constitute physical "
        "scale-out even after equalizing task capacity, and that computational "
        "benchmarking outputs should remain separate from prediction-time "
        "clinical features. The pipeline, configuration, aggregate results, and "
        "split diagnostics are publicly available.",
    )

    # ---------------------------------------------------------------
    # Table II: early-window row count 31,252 -> 29,886
    # ---------------------------------------------------------------
    t2 = tables[1]
    found = False
    for r in t2.rows:
        if r.cells[0].text.strip() == "Early-window clinical":
            assert r.cells[2].text.strip() == "31,252", r.cells[2].text
            r.cells[2].text = "29,886"
            found = True
    assert found, "Table II early-window row not found"

    # ---------------------------------------------------------------
    # III.D model paragraph: document the recalibration step actually used
    # ---------------------------------------------------------------
    fix_substring(
        paras[34],
        "Outcomes included AUROC, AUPRC, Brier score, sensitivity, specificity, "
        "and calibration intercept and slope. The 0.5 threshold was retained "
        "for descriptive threshold metrics; no clinical decision threshold was "
        "optimized.",
        "Because class-weighted training systematically shifts predicted "
        "probabilities away from the true prevalence, each model's raw output "
        "was recalibrated within each training fold using sigmoid (Platt) "
        "scaling fitted with an inner 5-fold, ungrouped cross-validation; the "
        "held-out test fold was never used at this step. Outcomes included "
        "AUROC, AUPRC, Brier score, Brier skill score (relative to a "
        "constant-prevalence baseline), sensitivity, specificity, and "
        "post-recalibration calibration intercept and slope. The 0.5 threshold "
        "was retained for descriptive threshold metrics; no clinical decision "
        "threshold was optimized.",
        "34-recal-method",
    )
    fix_substring(
        paras[35],
        "The prevalence-matched folds contained approximately 25,000 training "
        "and 6,250 test admissions, with 20,993–21,008 training patients and "
        "5,243–5,258 test patients. Calibration coefficients were estimated "
        "from predicted probabilities; because no recalibration model was "
        "fitted, they describe the native probability scale of each baseline.",
        "The prevalence-matched folds contained approximately 23,908 training "
        "and 5,978 test admissions, with 20,200–20,236 training patients and "
        "5,035–5,071 test patients. Calibration coefficients were estimated "
        "from predicted probabilities after the within-training-fold "
        "recalibration described above; they describe the recalibrated, not "
        "the native, probability scale of each baseline.",
        "35-splits",
    )

    # ---------------------------------------------------------------
    # III.E feature engineering: labs are count-only (fixed), not
    # mean/median/min/max/SD -- that description applies to vitals only
    # ---------------------------------------------------------------
    fix_substring(
        paras[37],
        "The early-window feature matrix combined demographic variables with "
        "distributional summaries of vital signs and laboratory observations "
        "available before the 48-h cutoff. For repeatedly measured variables, "
        "the pipeline derived count, mean, median, minimum, maximum, and "
        "standard deviation where applicable. An abnormal-laboratory count "
        "summarized the number of observations marked outside their reference "
        "range.",
        "The early-window feature matrix combined demographic variables with "
        "distributional summaries of vital signs and two laboratory-usage "
        "signals available before the 48-h cutoff. For each repeatedly measured "
        "vital sign, the pipeline derived count, mean, median, minimum, "
        "maximum, and standard deviation. Laboratory results were deliberately "
        "restricted to a count of observations and a count of values flagged "
        "outside their reference range, omitting mean/median/min/max/SD: "
        "LABEVENTS item identifiers span chemically unrelated tests (e.g., "
        "sodium, creatinine, hemoglobin, pH) with different units and reference "
        "ranges, so an itemid-unaware average of raw values has no clinical "
        "meaning. A correct fix requires joining item identifiers through "
        "D_LABITEMS, which was not available in the local extract; only the "
        "two itemid-agnostic usage signals and their missingness indicators "
        "were therefore retained.",
        "37-feature-eng",
    )

    # ---------------------------------------------------------------
    # Insert the local[6] vs standalone2 equal-capacity control result,
    # right before Fig. 2, at the end of Section IV.A.
    # ---------------------------------------------------------------
    anchor = paras[49]
    assert anchor.text.startswith("Fig. 2."), anchor.text[:60]
    new_p = anchor.insert_paragraph_before(
        "Because standalone configurations declared six total executor cores "
        "while local[8] exposed eight task threads, we ran a supplementary, "
        "same-session, block-randomized comparison of an equal-capacity "
        "local[6] configuration (six task threads in one JVM, matched to "
        "standalone's executor-core budget) against the two-worker standalone "
        "arrangement (standalone2). Local[6] and standalone2 were interleaved "
        "within each of 12 paired blocks (8 CPU/8 GiB per block) so that "
        "time-varying host load affected both arms equally. Local[6] remained "
        "faster than standalone2 for both workloads: admission-level, 238.6 ± "
        "13.4 s vs. 278.8 ± 67.9 s (paired difference −40.2 s, 95% CI [−80.9, "
        "0.5], geometric ratio 0.87, paired t = −2.17, p = 0.053); six-hour "
        "window, 278.8 ± 73.0 s vs. 319.7 ± 81.1 s (difference −40.9 s, 95% CI "
        "[−51.5, −30.3], ratio 0.87, paired t = −8.50, p < 0.0001). The direction "
        "was consistent across workloads, reaching significance for the "
        "six-hour window and remaining marginal for the admission-level "
        "workload at this sample size, indicating that the local[8] advantage "
        "is not attributable solely to its larger task-slot count.",
        style=anchor.style,
    )
    anchor._p.addprevious(new_p._p)

    # ---------------------------------------------------------------
    # Results IV.B: cohort size, prevalence, AUROC/AUPRC/Brier/calibration
    # ---------------------------------------------------------------
    assert paras[53].text.startswith("The early-window cohort contained 31,252"), paras[53].text[:80]
    assert paras[54].text.startswith("Repeated grouped holdouts supported"), paras[54].text[:80]
    replace_paragraph_text(
        paras[53],
        "The early-window cohort contained 29,886 admissions from 25,271 "
        "patients; mortality prevalence was 12.39%. This cohort differs from a "
        "preliminary version that defined the observation window over the "
        "multi-stay min/max envelope of each admission rather than its first "
        "ICU stay, which misclassified 534/31,252 admissions (1.7%) by "
        "including inter-stay service gaps in the eligibility window; the "
        "landmark was redefined around each admission's first icustay_id "
        "before final analysis. All five cross-validation folds had disjoint "
        "patient groups and closely matched outcome prevalence. Histogram "
        "gradient boosting had the highest discrimination: AUROC 0.815 ± 0.005 "
        "and AUPRC 0.402 ± 0.025 in five-fold grouped cross-validation. "
        "Logistic regression yielded AUROC 0.770 ± 0.014 and AUPRC 0.341 ± "
        "0.034; random forest yielded AUROC 0.810 ± 0.007 and AUPRC 0.386 ± "
        "0.020. Each model's raw class-weighted probabilities were recalibrated "
        "within each training fold (sigmoid/Platt scaling, 5-fold inner "
        "cross-validation, never touching the corresponding test fold) before "
        "scoring. After recalibration, gradient boosting's Brier score was "
        "0.090 ± 0.001 — below the constant-prevalence baseline of 0.109, for a "
        "Brier skill score of 0.169 ± 0.013 — and its calibration slope and "
        "intercept were 1.065 ± 0.039 and 0.092 ± 0.054, close to the ideal "
        "values of 1 and 0. Without this recalibration step the raw "
        "class-weighted probabilities carried a large negative intercept, "
        "indicating systematic miscalibration despite useful ranking; "
        "recalibration corrects the probability level without altering rank "
        "order or AUROC/AUPRC.",
    )
    replace_paragraph_text(
        paras[54],
        "Repeated grouped holdouts supported the cross-validation ordering. "
        "Gradient boosting achieved mean AUROC 0.815 (range 0.800–0.830) and "
        "mean AUPRC 0.402 (0.371–0.439); random forest achieved 0.811 "
        "(0.796–0.828) and 0.386 (0.351–0.418); logistic regression achieved "
        "0.769 (0.749–0.787) and 0.337 (0.310–0.370). Threshold-dependent "
        "results were consistent across all three recalibrated models: at the "
        "unoptimized 0.5 threshold, sensitivity was low and specificity was "
        "high for gradient boosting (0.100 and 0.992), random forest (0.137 and "
        "0.985), and logistic regression (0.076 and 0.991), while precision "
        "remained comparatively high (0.55–0.64) because false positives were "
        "rare. This pattern is the expected consequence of properly calibrated "
        "probabilities under 12.4% prevalence — few predictions exceed 0.5 for "
        "a minority outcome even when ranking is informative — and illustrates "
        "why AUROC alone is insufficient and why the unoptimized threshold must "
        "not be construed as a clinical operating point.",
    )

    # ---------------------------------------------------------------
    # Missingness / feature-importance: drop the now-inapplicable
    # "laboratory extrema and dispersion" phrase (labs are count-only)
    # ---------------------------------------------------------------
    fix_substring(
        paras[55],
        "Random-forest importance ranked age, abnormal laboratory count, "
        "laboratory extrema and dispersion, respiratory rate, temperature, "
        "oxygen saturation, and mean arterial pressure among the leading "
        "predictors.",
        "Random-forest importance ranked the abnormal-laboratory count, age, "
        "the laboratory-observation count, respiratory rate, temperature, "
        "oxygen saturation, and mean arterial pressure among the leading "
        "predictors.",
        "55-feature-importance",
    )

    # ---------------------------------------------------------------
    # IV.C robustness section
    # ---------------------------------------------------------------
    assert paras[60].text.startswith("The primary and repeated-holdout estimates"), paras[60].text[:80]
    replace_paragraph_text(
        paras[60],
        "The primary and repeated-holdout estimates were closely aligned. For "
        "gradient boosting, the mean AUROC was 0.815 in both analyses (cross-"
        "validation 0.8151 vs. repeated holdout 0.8152), while mean AUPRC "
        "differed by 0.0001. Corresponding random-forest estimates differed by "
        "0.0010 for AUROC and 0.0001 for AUPRC. This agreement indicates "
        "limited sensitivity to the particular resampling scheme used here, "
        "but it is not a substitute for temporal or external validation "
        "because every split originates from the same database and era.",
    )
    assert paras[61].text.startswith("The Brier score and calibration"), paras[61].text[:80]
    replace_paragraph_text(
        paras[61],
        "The Brier score and calibration coefficients require joint "
        "interpretation even after recalibration. Gradient boosting had both "
        "the lowest Brier score and the calibration slope closest to one, "
        "consistent with its higher discrimination; random forest and logistic "
        "regression had modestly higher Brier scores and slopes further from "
        "one (0.092/1.046 and 0.096/1.007, calibration intercepts 0.071 and "
        "0.008). All three intercepts were close to zero, indicating that "
        "within-training-fold recalibration removed the systematic level shift "
        "present in the raw class-weighted probabilities. A model can still "
        "rank admissions well while carrying residual slope miscalibration; "
        "reporting slope and intercept alongside AUROC and Brier score makes "
        "this distinction visible rather than assuming recalibration alone "
        "guarantees a decision-ready probability.",
    )
    fix_substring(
        paras[62],
        "Because mortality prevalence was 12.58%, AUPRC provides a more "
        "informative view of positive-case retrieval than accuracy. The "
        "observed AUPRC values substantially exceeded the no-skill prevalence "
        "reference, yet precision at the fixed 0.5 threshold remained modest "
        "for gradient boosting and logistic regression.",
        "Because mortality prevalence was 12.39%, AUPRC provides a more "
        "informative view of positive-case retrieval than accuracy. The "
        "observed AUPRC values substantially exceeded the no-skill prevalence "
        "reference, and precision at the fixed 0.5 threshold was in fact "
        "comparatively high (0.53–0.64) precisely because so few admissions "
        "crossed the threshold; this reflects low recall rather than a "
        "favorable operating point.",
        "62-auprc-precision",
    )

    # ---------------------------------------------------------------
    # Discussion: note the equal-capacity control's result
    # ---------------------------------------------------------------
    fix_substring(
        paras[64],
        "The observed result is an implemented-topology comparison, not a "
        "causal estimate of execution mode alone, because task concurrency, "
        "Java heaps, process count, and driver workload were not perfectly "
        "equalized.",
        "The observed result is an implemented-topology comparison, not a "
        "causal estimate of execution mode alone, because task concurrency, "
        "Java heaps, process count, and driver workload were not perfectly "
        "equalized; a supplementary equal-task-capacity comparison (local[6] "
        "vs. two-worker standalone) found the same direction, suggesting the "
        "effect is not solely a task-count artifact.",
        "64-discussion",
    )

    # ---------------------------------------------------------------
    # Section VII: replace "future factorial study" framing with what the
    # local[6] control already showed
    # ---------------------------------------------------------------
    fix_substring(
        paras[79],
        "Executor cores were fixed at six in the standalone arrangements while "
        "local[8] exposed eight task threads; this is part of the tested "
        "topology and one reason the experiment cannot isolate a pure "
        "communication penalty. A future factorial study should vary process "
        "model and task capacity independently, record processor frequency and "
        "context switches, and repeat the experiment on physically separate "
        "workers.",
        "Executor cores were fixed at six in the standalone arrangements while "
        "local[8] exposed eight task threads; Section IV.A reports a "
        "supplementary equal-capacity comparison (local[6] vs. standalone2) "
        "that partially disentangles this confound and still favors local "
        "execution, though the admission-level contrast was only marginal at "
        "n = 12 paired blocks and the comparison shares the same host and "
        "session as the rest of the study. It therefore cannot isolate a pure "
        "communication penalty from other process-boundary costs. A future "
        "factorial study should vary process model and task capacity "
        "independently across a larger number of blocks, record processor "
        "frequency and context switches, and repeat the experiment on "
        "physically separate workers and on independent days.",
        "79-sec7",
    )

    # ---------------------------------------------------------------
    # Data and Code Availability: commit hash
    # ---------------------------------------------------------------
    fix_substring(
        paras[91],
        "commit 55c8926aa6d1d384dc3587aab904d0bed02b3d52",
        "commit d0e87f9696562554185b5c7949daf006ffcc75fb",
        "91-data-avail",
    )

    d.save(DOC)

    # ---------------------------------------------------------------
    # Table VI: Model | AUROC | AUPRC | Brier | Cal. slope
    # ---------------------------------------------------------------
    d = Document(DOC)  # reload so table objects are consistent post-save
    tables = d.tables
    t6 = tables[6]
    header = [c.text.strip() for c in t6.rows[0].cells]
    assert header == ["Model", "AUROC", "AUPRC", "Brier", "Cal. slope"], header
    data6 = {
        "Logistic regression": ("0.770 ± .014", "0.341 ± .034", "0.096 ± .002", "1.007 ± .100"),
        "Random forest":       ("0.810 ± .007", "0.386 ± .020", "0.092 ± .002", "1.046 ± .032"),
        "Gradient boosting":   ("0.815 ± .005", "0.402 ± .025", "0.090 ± .001", "1.065 ± .039"),
    }
    for row in t6.rows[1:]:
        model = row.cells[0].text.strip()
        assert model in data6, model
        auroc, auprc, brier, slope = data6[model]
        row.cells[1].text = auroc
        row.cells[2].text = auprc
        row.cells[3].text = brier
        row.cells[4].text = slope

    # ---------------------------------------------------------------
    # Table VII: Model | AUROC range | AUPRC range | Sensitivity | Specificity
    # ---------------------------------------------------------------
    t7 = tables[7]
    header7 = [c.text.strip() for c in t7.rows[0].cells]
    assert header7 == ["Model", "AUROC range", "AUPRC range", "Sensitivity", "Specificity"], header7
    data7 = {
        "Gradient Boosting":   ("0.800–0.830", "0.371–0.439", "0.100", "0.992"),
        "Logistic Regression": ("0.749–0.787", "0.310–0.370", "0.076", "0.991"),
        "Random Forest":       ("0.796–0.828", "0.351–0.418", "0.137", "0.985"),
    }
    for row in t7.rows[1:]:
        model = row.cells[0].text.strip()
        assert model in data7, model
        ar, apr, sens, spec = data7[model]
        row.cells[1].text = ar
        row.cells[2].text = apr
        row.cells[3].text = sens
        row.cells[4].text = spec

    d.save(DOC)
    print(f"Saved: {DOC}")


if __name__ == "__main__":
    main()
