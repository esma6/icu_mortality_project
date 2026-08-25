"""Second round of fixes to IEEE_JBHI_Submission_Manuscript.docx, from an
independent (Codex) review of the round-1 rebuild plus a self-caught
Methods/Results inconsistency found while re-reading the result.

Findings addressed:
1. CalibratedClassifierCV(pipeline, method="sigmoid", cv=5) used an ungrouped
   inner 5-fold split for calibration, so a patient's admissions could span
   the inner model-fitting and calibration folds (never the outer test fold,
   so outer discrimination metrics were unaffected, but the "leakage-aware"
   framing did not fully extend to the calibration step). Fixed in
   scripts/train_leakfree_model.py to use StratifiedGroupKFold for the inner
   split too, and scripts/train_leakfree_model.py was re-run, changing every
   AUROC/AUPRC/Brier/BSS/calibration number in this manuscript by a small
   (2nd-3rd decimal) amount. Cohort size, prevalence, split diagnostics,
   missingness, feature importance, and all ETL/local[6] numbers are
   unaffected (verified unchanged in outputs/tables_ml_leakfree/*.csv).
2. Section III.D described the inner calibration CV as "ungrouped" -- true
   before fix 1, false after; a separate self-caught inconsistency (the
   Results section had already been updated to say "grouped" but Methods had
   not).
3. Section VII's limitations paragraph claimed the study "lacks... post-hoc
   probability recalibration" -- directly contradicted by the recalibration
   methodology already central to the manuscript.
4. No Funding / Conflict of Interest sections (both required by IEEE JBHI;
   author-confirmed: no external funding, no conflict of interest).
5. Author block was a single free-text line, not IEEE's numbered-footnote
   affiliation format.
6. Abstract's local[6] sentence could be skimmed as claiming both workload
   contrasts were significant; the p-values were already both reported but
   the significant/non-significant distinction was only explicit in the
   body text, not the abstract.

Run after build_ieee_jbhi_manuscript.py (round 1) and after re-running
scripts/train_leakfree_model.py.
"""
import sys
sys.stdout.reconfigure(encoding="utf-8")
from docx import Document

DOC = "IEEE_JBHI_Submission_Manuscript.docx"


def replace_paragraph_text(paragraph, new_text):
    for run in list(paragraph.runs)[1:]:
        run.text = ""
    if paragraph.runs:
        paragraph.runs[0].text = new_text
    else:
        paragraph.add_run(new_text)


def fix_substring(paragraph, old, new, label):
    text = paragraph.text
    if old not in text:
        raise ValueError(f"[{label}] not found: {old!r}\nIn: {text[:200]!r}")
    replace_paragraph_text(paragraph, text.replace(old, new))


def main():
    d = Document(DOC)
    paras = d.paragraphs

    # -------------------------------------------------------------
    # Author block -> IEEE-style numbered affiliation footnotes
    # -------------------------------------------------------------
    assert paras[1].text == "Esma Fazilet Karagülle and Esra Odabaş Yıldırım", paras[1].text
    replace_paragraph_text(paras[1], "Esma Fazilet Karagülle¹ and Esra Odabaş Yıldırım²")
    assert paras[2].text.startswith("Atatürk University, Erzurum"), paras[2].text
    replace_paragraph_text(
        paras[2],
        "¹Department of Computer Engineering, Atatürk University, Erzurum, "
        "Türkiye (e-mail: esmafazilet.karagulle17@ogr.atauni.edu.tr; corresponding "
        "author). ²Department of Software Engineering, Atatürk University, "
        "Erzurum, Türkiye (e-mail: esra.odabas@atauni.edu.tr).",
    )

    # -------------------------------------------------------------
    # Abstract: clarity + updated ML numbers (post grouped-calibration fix)
    # -------------------------------------------------------------
    fix_substring(
        paras[3],
        "a supplementary equal-capacity comparison (local[6] vs. a two-worker "
        "standalone arrangement) was run in the same session; local[6] remained "
        "faster for both workloads (six-hour window: p < 0.0001; admission-level: "
        "p = 0.053).",
        "a supplementary equal-capacity comparison (local[6] vs. a two-worker "
        "standalone arrangement) was run in the same session; local[6] was faster "
        "in both workloads, reaching significance for the six-hour window (p < "
        "0.0001) but not for the admission-level workload (p = 0.053).",
        "abstract-local6-clarity",
    )
    fix_substring(
        paras[3],
        "Gradient boosting achieved an area under the "
        "receiver-operating-characteristic curve of 0.815 ± 0.005 and an area "
        "under the precision-recall curve of 0.402 ± 0.025 in five-fold "
        "patient-grouped cross-validation; within-training-fold recalibration "
        "of the class-weighted probabilities corrected a large negative "
        "calibration intercept to 0.092 ± 0.054 and yielded a Brier skill "
        "score of 0.169 relative to a constant-prevalence baseline.",
        "Gradient boosting achieved an area under the "
        "receiver-operating-characteristic curve of 0.813 ± 0.006 and an area "
        "under the precision-recall curve of 0.402 ± 0.023 in five-fold "
        "patient-grouped cross-validation; within-training-fold, patient-grouped "
        "recalibration of the class-weighted probabilities corrected a large "
        "negative calibration intercept to 0.099 ± 0.053 and yielded a Brier "
        "skill score of 0.168 relative to a constant-prevalence baseline.",
        "abstract-ml-numbers",
    )

    # -------------------------------------------------------------
    # Methods III.D: describe the inner calibration CV as grouped
    # -------------------------------------------------------------
    fix_substring(
        paras[34],
        "each model's raw output was recalibrated within each training fold "
        "using sigmoid (Platt) scaling fitted with an inner 5-fold, ungrouped "
        "cross-validation; the held-out test fold was never used at this step.",
        "each model's raw output was recalibrated within each training fold "
        "using sigmoid (Platt) scaling fitted with an inner 5-fold "
        "cross-validation itself grouped by subject_id (StratifiedGroupKFold), "
        "so a patient's admissions cannot span the inner model-fitting and "
        "calibration folds either; the held-out outer test fold was never used "
        "at this step.",
        "34-grouped-inner-cv",
    )

    # -------------------------------------------------------------
    # B. Leakage-Aware Clinical Validation: cohort/CV results paragraph
    # -------------------------------------------------------------
    assert paras[54].text.startswith("The early-window cohort contained 29,886"), paras[54].text[:80]
    replace_paragraph_text(
        paras[54],
        "The early-window cohort contained 29,886 admissions from 25,271 "
        "patients; mortality prevalence was 12.39%. This cohort differs from a "
        "preliminary version that defined the observation window over the "
        "multi-stay min/max envelope of each admission rather than its first "
        "ICU stay, which misclassified 534/31,252 admissions (1.7%) by "
        "including inter-stay service gaps in the eligibility window; the "
        "landmark was redefined around each admission's first icustay_id "
        "before final analysis. All five cross-validation folds had disjoint "
        "patient groups and closely matched outcome prevalence. Histogram "
        "gradient boosting had the highest discrimination: AUROC 0.813 ± 0.006 "
        "and AUPRC 0.402 ± 0.023 in five-fold grouped cross-validation. "
        "Logistic regression yielded AUROC 0.770 ± 0.014 and AUPRC 0.341 ± "
        "0.034; random forest yielded AUROC 0.809 ± 0.006 and AUPRC 0.386 ± "
        "0.021. Each model's raw class-weighted probabilities were recalibrated "
        "within each training fold using an inner 5-fold split itself grouped "
        "by subject_id (sigmoid/Platt scaling; StratifiedGroupKFold; never "
        "touching the corresponding outer test fold) before scoring. After "
        "recalibration, gradient boosting's Brier score was 0.090 ± 0.001 — "
        "below the constant-prevalence baseline of 0.109, for a Brier skill "
        "score of 0.168 ± 0.012 — and its calibration slope and intercept were "
        "1.068 ± 0.035 and 0.099 ± 0.053, close to the ideal values of 1 and 0. "
        "Without this recalibration step the raw class-weighted probabilities "
        "carried a large negative intercept, indicating systematic "
        "miscalibration despite useful ranking; recalibration corrects the "
        "probability level without altering rank order or AUROC/AUPRC.",
    )
    assert paras[55].text.startswith("Repeated grouped holdouts supported"), paras[55].text[:80]
    replace_paragraph_text(
        paras[55],
        "Repeated grouped holdouts supported the cross-validation ordering. "
        "Gradient boosting achieved mean AUROC 0.816 (range 0.800–0.833) and "
        "mean AUPRC 0.404 (0.375–0.445); random forest achieved 0.810 "
        "(0.795–0.828) and 0.386 (0.355–0.416); logistic regression achieved "
        "0.768 (0.749–0.787) and 0.337 (0.311–0.370). Threshold-dependent "
        "results were consistent across all three recalibrated models: at the "
        "unoptimized 0.5 threshold, sensitivity was low and specificity was "
        "high for gradient boosting (0.098 and 0.992), random forest (0.137 and "
        "0.985), and logistic regression (0.074 and 0.991), while precision "
        "remained comparatively high (0.55–0.63) because false positives were "
        "rare. This pattern is the expected consequence of properly calibrated "
        "probabilities under 12.4% prevalence — few predictions exceed 0.5 for "
        "a minority outcome even when ranking is informative — and illustrates "
        "why AUROC alone is insufficient and why the unoptimized threshold must "
        "not be construed as a clinical operating point.",
    )

    # -------------------------------------------------------------
    # C. Robustness section
    # -------------------------------------------------------------
    assert paras[61].text.startswith("The primary and repeated-holdout estimates"), paras[61].text[:80]
    replace_paragraph_text(
        paras[61],
        "The primary and repeated-holdout estimates were reasonably aligned, "
        "though not identical, given the different resampling schemes. For "
        "gradient boosting, mean AUROC was 0.813 in five-fold cross-validation "
        "vs. 0.816 in repeated holdout (difference 0.003), while mean AUPRC "
        "differed by 0.002 (0.402 vs. 0.404). Corresponding random-forest "
        "estimates differed by 0.001 for AUROC (0.809 vs. 0.810) and were "
        "identical to three decimal places for AUPRC (0.386). This agreement "
        "indicates limited sensitivity to the particular resampling scheme "
        "used here, but it is not a substitute for temporal or external "
        "validation because every split originates from the same database and "
        "era.",
    )
    assert paras[62].text.startswith("The Brier score and calibration"), paras[62].text[:80]
    replace_paragraph_text(
        paras[62],
        "The Brier score and calibration coefficients require joint "
        "interpretation even after recalibration. Gradient boosting had the "
        "lowest Brier score (0.090), consistent with its higher discrimination, "
        "while logistic regression's calibration slope was closest to one "
        "(1.010, vs. 1.068 for gradient boosting and 1.046 for random forest). "
        "All three calibration intercepts were close to zero (0.014–0.099), "
        "indicating that within-training-fold, patient-grouped recalibration "
        "removed the systematic level shift present in the raw class-weighted "
        "probabilities. A model can still rank admissions well while carrying "
        "residual slope miscalibration; reporting slope and intercept alongside "
        "AUROC and Brier score makes this distinction visible rather than "
        "assuming recalibration alone guarantees a decision-ready probability.",
    )
    fix_substring(
        paras[63],
        "precision at the fixed 0.5 threshold was in fact comparatively high "
        "(0.53–0.64) precisely",
        "precision at the fixed 0.5 threshold was in fact comparatively high "
        "(0.54–0.64) precisely",
        "63-precision-range",
    )

    # -------------------------------------------------------------
    # Section VII limitations: recalibration is no longer an absent item
    # -------------------------------------------------------------
    fix_substring(
        paras[71],
        "The clinical validation lacks external validation, uncertainty "
        "intervals based on independent sites, decision-curve analysis, and "
        "post-hoc probability recalibration.",
        "Probabilities are recalibrated within each training fold (Section "
        "III.D), but the clinical validation still lacks external validation, "
        "uncertainty intervals based on independent sites, and decision-curve "
        "analysis.",
        "71-limitations-recal",
    )

    d.save(DOC)

    # -------------------------------------------------------------
    # Funding / Conflict of Interest sections (inserted before Ethics)
    # -------------------------------------------------------------
    d = Document(DOC)
    paras = d.paragraphs
    anchor = None
    for p in paras:
        if p.text == "ETHICS AND DATA GOVERNANCE":
            anchor = p
            break
    assert anchor is not None, "ETHICS AND DATA GOVERNANCE heading not found"
    heading_style = None
    body_style = None
    for i, p in enumerate(paras):
        if p.text == "ACKNOWLEDGMENT":
            heading_style = p.style
            body_style = paras[i + 1].style
            break
    assert heading_style is not None and body_style is not None

    anchor.insert_paragraph_before("FUNDING", style=heading_style)
    anchor.insert_paragraph_before("This research received no external funding.", style=body_style)
    anchor.insert_paragraph_before("CONFLICT OF INTEREST", style=heading_style)
    anchor.insert_paragraph_before("The authors declare no conflict of interest.", style=body_style)

    d.save(DOC)

    # -------------------------------------------------------------
    # Table VI / VII: updated post-grouped-calibration numbers
    # -------------------------------------------------------------
    d = Document(DOC)
    tables = d.tables
    t6 = tables[6]
    header = [c.text.strip() for c in t6.rows[0].cells]
    assert header == ["Model", "AUROC", "AUPRC", "Brier", "Cal. slope"], header
    data6 = {
        "Logistic regression": ("0.770 ± .014", "0.341 ± .034", "0.096 ± .002", "1.010 ± .097"),
        "Random forest":       ("0.809 ± .006", "0.386 ± .021", "0.092 ± .002", "1.046 ± .027"),
        "Gradient boosting":   ("0.813 ± .006", "0.402 ± .023", "0.090 ± .001", "1.068 ± .035"),
    }
    for row in t6.rows[1:]:
        model = row.cells[0].text.strip()
        assert model in data6, model
        auroc, auprc, brier, slope = data6[model]
        row.cells[1].text = auroc
        row.cells[2].text = auprc
        row.cells[3].text = brier
        row.cells[4].text = slope

    t7 = tables[7]
    header7 = [c.text.strip() for c in t7.rows[0].cells]
    assert header7 == ["Model", "AUROC range", "AUPRC range", "Sensitivity", "Specificity"], header7
    data7 = {
        "Gradient Boosting":   ("0.800–0.833", "0.375–0.445", "0.098", "0.992"),
        "Logistic Regression": ("0.749–0.787", "0.311–0.370", "0.074", "0.991"),
        "Random Forest":       ("0.795–0.828", "0.355–0.416", "0.137", "0.985"),
    }
    for row in t7.rows[1:]:
        model = row.cells[0].text.strip()
        assert model in data7, model
        ar, apr, sens, spec = data7[model]
        row.cells[1].text = ar
        row.cells[2].text = apr
        row.cells[3].text = sens
        row.cells[4].text = spec

    d.save(DOC)
    print(f"Saved: {DOC}")


if __name__ == "__main__":
    main()
