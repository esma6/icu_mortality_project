from __future__ import annotations

import re
from pathlib import Path

import pandas as pd
from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
SRC = Path(__file__).resolve().parent / "MIMICIII_Spark_ETL_Makale_Literatur_Temelli_v9.docx"
DST = Path(__file__).resolve().parent / "MIMICIII_Spark_ETL_Makale_Literatur_Temelli_v10.docx"
ML_TABLES = ROOT / "outputs" / "tables_ml_leakfree"
ML_FIGURES = ROOT / "outputs" / "figures_ml_leakfree"

HEADER_FILL = "1F4E79"
ZEBRA_FILL = "EAF2F8"


def tr(x: float, decimals: int = 1) -> str:
    return f"{x:.{decimals}f}".replace(".", ",")


def shade_cell(cell, fill: str) -> None:
    tcPr = cell._tc.get_or_add_tcPr()
    shd = tcPr.makeelement(qn("w:shd"), {qn("w:val"): "clear", qn("w:color"): "auto", qn("w:fill"): fill})
    tcPr.append(shd)


def set_cell_text(cell, text: str, bold: bool = False, white: bool = False, size: int = 8) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.name = "Arial"
    run.font.bold = bold
    if white:
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    p.paragraph_format.space_after = Pt(0)


def build_table(document: Document, anchor_p, headers: list[str], rows: list[list[str]]):
    n_cols = len(headers)
    table = document.add_table(rows=1, cols=n_cols)
    table.style = document.styles["Table Grid"]
    for j, h in enumerate(headers):
        cell = table.rows[0].cells[j]
        set_cell_text(cell, h, bold=True, white=True)
        shade_cell(cell, HEADER_FILL)
    for i, row in enumerate(rows):
        tr_cells = table.add_row().cells
        for j, val in enumerate(row):
            set_cell_text(tr_cells[j], val, bold=False, white=False)
            if i % 2 == 0:
                shade_cell(tr_cells[j], ZEBRA_FILL)
    anchor_p._p.addprevious(table._tbl)
    return table


def replace_paragraph_text(paragraph, new_text: str) -> None:
    for run in list(paragraph.runs):
        run.text = ""
    if paragraph.runs:
        paragraph.runs[0].text = new_text
    else:
        paragraph.add_run(new_text)


def main() -> None:
    document = Document(SRC)
    paras = document.paragraphs

    # Capture anchors BEFORE any structural insertion (indices are stable — python-docx
    # Paragraph objects wrap a specific <w:p> element; inserting elsewhere in the tree
    # does not invalidate these references).
    p_abstract_tr = paras[4]
    p_abstract_en = paras[7]
    p_contrib_intro = paras[24]
    p_contrib_b1 = paras[25]
    p_contrib_b4 = paras[28]
    p_37_anchor = paras[64]     # heading "4. Bulgular" -> insert 3.7 right before it
    p_46_anchor = paras[101]    # heading "5. Tartışma" -> insert 4.6 right before it
    p_53_b1 = paras[111]        # 5.3 first body paragraph
    p_53_b2 = paras[112]        # 5.3 second body paragraph (future-work list, to rewrite)
    p_64_body = paras[124]      # 6.4 body
    p_65_body = paras[126]      # 6.5 body
    p_conclusion = paras[128]   # 7. Sonuç body
    p_checklist_ml = paras[172] # checklist bullet about future clinical mortality model

    # ---------- Abstract rewrites ----------
    replace_paragraph_text(
        p_abstract_tr,
        "Arka plan: Klinik EHR veri hatları 2020 sonrasında standartlaştırılmış zaman "
        "serilerinden olay-merkezli ve kurumlar arası taşınabilir temsillere ilerlemiştir. "
        "Kaynak kısıtlı ortamlarda özellik granülerliği ile Spark yürütme topolojisinin "
        "birlikte etkisi yeterince ölçülmemiş, üretilen özellik uzaylarının klinik model "
        "doğrulamasındaki kullanılabilirliği de aynı çerçevede gösterilmemiştir. Amaç: Aynı "
        "fiziksel kaynak bütçesinde Spark local iş parçacığı paralelliği ile eş-konumlu "
        "standalone worker topolojilerini iki MIMIC-III özellik rejiminde karşılaştırmak ve "
        "ayrı, sızıntısız bir erken-pencere özellik setiyle klinik model doğrulamasını aynı "
        "deneysel hatta göstermek. Yöntem: Docker Compose ile kapsayıcılaştırılmış Spark "
        "3.5.1 hattı, 12 mantıksal işlemci ve tek NVMe SSD'li aynı Windows 11 hostunda "
        "local[2], local[4], local[8], master+bir worker ve master+iki worker olarak "
        "çalıştırıldı. Compact ürün 58.976, timeseries ürün 1.180.395 satır üretti; "
        "birincil seride n=5 ve randomize, kaynak-kotalı doğrulama serisinde n=12 ölçümle "
        "analiz edildi. Ayrıca ICU kalışı 48 saatten kısa kabuller dışlanarak sızıntısız, "
        "hasta-düzeyinde gruplanmış bir erken-pencere kohortu (31.252 kabul, 26.251 hasta, "
        "%12,6 mortalite) StratifiedGroupKFold ile doğrulandı. Bulgular: En düşük ortalama "
        "süre compact için local[8]'de 204,8±2,9 s, timeseries için 263,9±30,1 s idi. "
        "Doğrulama serisi bu üstünlüğü her iki iş yükünde de Holm düzeltmeli testlerle "
        "doğruladı. Sızıntısız kohortta en iyi ayırt edicilik HistGradientBoosting ile elde "
        "edildi (AUROC 0,814±0,007, kalibrasyon eğimi 1,06±0,05). Sonuç: İncelenen tek-host "
        "ortamında worker kapsayıcısı eklemek fiziksel kapasite eklememiş, local[8]'i "
        "geçmemiştir. Bulgular gerçek çok-hostlu strong scaling kanıtı değil, sabit kaynak "
        "bütçesinde yürütme-topolojisi karakterizasyonudur. Aynı hattan türetilen sızıntısız "
        "özellik uzayı, ilk 48 saati atlatan hastalarda hastane-içi mortalite riskini "
        "klinik olarak makul bir ayırt edicilikle tahmin edebilmektedir.",
    )
    replace_paragraph_text(
        p_abstract_en,
        "Background: Since 2020, clinical EHR pipelines have progressed from standardized "
        "time-series extraction toward event-centric and interoperable representations. "
        "The joint effect of feature granularity and Spark execution topology remains "
        "under-characterized in resource-constrained environments, and the downstream "
        "usability of the resulting feature space for clinical model validation has not "
        "been shown within the same experimental framework. Objective: To compare local "
        "thread parallelism with co-located standalone Spark workers under a fixed "
        "physical resource budget, and to demonstrate clinical model validation on a "
        "separate, leak-free early-window feature set derived from the same pipeline. "
        "Methods: A Docker Compose-containerized Spark 3.5.1 ETL pipeline processed "
        "MIMIC-III v1.4 on one Windows 11 host with 12 logical processors and one NVMe "
        "SSD. The compact product yielded 58,976 rows and the timeseries product "
        "1,180,395 rows, analyzed with n=5 in the primary series and n=12 in an "
        "independent randomized, resource-quota-controlled validation series. Admissions "
        "with ICU stays shorter than 48 hours were excluded to build a leak-free, "
        "patient-grouped early-window cohort (31,252 admissions, 26,251 patients, 12.6% "
        "mortality), validated with StratifiedGroupKFold. Results: Local[8] had the "
        "lowest mean runtime for compact (204.8±2.9 s) and timeseries (263.9±30.1 s). The "
        "validation series confirmed this advantage for both workloads with Holm-adjusted "
        "contrasts. In the leak-free cohort, HistGradientBoosting achieved the best "
        "discrimination (AUROC 0.814±0.007, calibration slope 1.06±0.05). Conclusion: On "
        "this single host, adding worker containers did not add physical capacity and did "
        "not outperform local[8]. This is a fixed-budget topology characterization, not "
        "evidence of multi-host strong scaling. The leak-free feature space derived from "
        "the same pipeline predicts in-hospital mortality risk among patients surviving "
        "the first 48 hours in ICU with clinically plausible discrimination.",
    )

    # ---------- Contributions ----------
    replace_paragraph_text(
        p_contrib_intro,
        "Bu çalışma yeni bir klinik veri standardı önermemektedir. Standart "
        "sınıflandırıcılar yeni bir model mimarisi olarak değil, üretilen özellik "
        "uzayının sızıntısız kullanılabilirliğini göstermek amacıyla eğitilmiştir. "
        "Katkıları daha dar ve sınanabilir beş noktadadır:",
    )
    new_contrib = p_contrib_b4.insert_paragraph_before(
        "üretilen özellik uzayının, ICU kalışı 48 saati aşan hastalarda hastane-içi "
        "mortalite riskini hasta-düzeyinde gruplanmış, sızıntısız bir doğrulama "
        "tasarımıyla klinik olarak makul bir ayırt edicilikle tahmin edebildiğini "
        "göstermesi.",
        style="List Bullet",
    )
    # move it to be the 5th (last) bullet, right after the current 4th
    p_contrib_b4._p.addnext(new_contrib._p)

    # ---------- 3.7 Methods subsection ----------
    p_37_anchor.insert_paragraph_before(
        "3.7. Sızıntısız erken pencere özellik seti ve model doğrulama tasarımı", style="Heading 2",
    )
    p_37_anchor.insert_paragraph_before(
        "Bölüm 5.3'te belirtilen sızıntı riskini gidermek için ayrı, amaca özel bir "
        "üçüncü özellik ürünü (early_window) oluşturulmuştur. Bu ürün compact ve "
        "timeseries ETL-karşılaştırma ürünlerinden bağımsızdır ve onları etkilemez, "
        "yalnızca Bölüm 4.6'daki klinik model doğrulaması için kullanılmıştır. "
        "Harutyunyan ve ark.'nın MIMIC hastane-içi mortalite kıyaslama kuralına "
        "uyularak [7], ICU kalışı 48 saatten kısa olan kabuller dışlanmış, kalan her "
        "kabul için yalnızca ICU girişinden (min intime) sonraki ilk 48 saat içindeki "
        "ölçümler kullanılmıştır. Laboratuvar sonuçlarına da ilk kez zaman damgası "
        "(charttime) eklenerek aynı 48 saatlik pencereyle sınırlandırılmıştır — önceki "
        "üç rejimde LABEVENTS zaman bilgisi taşımadığından tüm yatış boyunca "
        "özetleniyordu. ICU çıkış zamanına veya toplam yatış süresine dayalı hiçbir "
        "özellik (icu_los_mean, icu_los_total, icu_stay_count) bu üründe "
        "hesaplanmamıştır, ilgili fonksiyon bu dalda hiç çağrılmamıştır.",
        style="Normal",
    )
    p_37_anchor.insert_paragraph_before(
        "Elde edilen kohort 58.976 kabulün 31.252'sini (%53,0) kapsar ve 26.251 "
        "benzersiz hastaya aittir, mortalite prevalansı %12,6'dır. Aynı hastanın "
        "birden fazla uygun yatışı olabildiğinden (31.252 kabul / 26.251 hasta), "
        "eğitim-test bölünmesi kabul (hadm_id) değil hasta (subject_id) düzeyinde "
        "yapılmıştır: scikit-learn StratifiedGroupKFold ile hem sınıf oranı yaklaşık "
        "korunmuş hem de aynı hastanın kayıtlarının eğitim ve test kümesine dağılması "
        "engellenmiştir. Tekrarlı ayrım (n=30 tekrar, test oranı %20) her tekrarda "
        "StratifiedGroupKFold(n_splits=5)'in ilk katmanı test kümesi olarak alınarak "
        "üretilmiş, ayrıca 5-katlı gruplu çapraz doğrulama ile tamamlanmıştır. Her "
        "bölünmede eğitim ve test hasta kümelerinin ayrık olduğu programatik olarak "
        "doğrulanmıştır (assert), split_diagnostics tabloları hasta/satır sayısı ve "
        "pozitif oranın tekrarlar arasında kararlı kaldığını göstermektedir (eğitim/"
        "test pozitif oranı yaklaşık %12,58-%12,59).",
        style="Normal",
    )
    p_37_anchor.insert_paragraph_before(
        "Üç sınıflandırıcı (Lojistik Regresyon, Rastgele Orman, HistGradientBoosting, "
        "scikit-learn varsayılan sınıf-dengeleme ağırlıklarıyla) medyan-imputasyon, "
        "%1-%99 kuantil kırpma ve standartlaştırmadan oluşan yalnızca-eğitim-kümesi bir "
        "ön işleme hattıyla eğitilmiştir. Değerlendirme AUROC, AUPRC, Brier skoru ve "
        "kalibrasyon eğrisi/intercept/slope'u (logit(olasılık) üzerinde lojistik "
        "regresyonla) içerir. Karar eğrisi analizi ve klinik alt grup performansı bu "
        "çalışmanın kapsamı dışında bırakılmış, gelecek çalışma olarak korunmuştur "
        "(Bölüm 6.5). Kohort tanımı ICU kalışının ilk 48 saatini aşan hastalarla "
        "sınırlı olduğundan bulgular 'ilk 48 saati atlatan hastalarda hastane-içi "
        "mortalite riski' için geçerlidir, genel ICU mortalite tahmini için değildir.",
        style="Normal",
    )

    # ---------- 4.6 Results subsection ----------
    cohort_flow = pd.read_csv(ML_TABLES / "cohort_flow.csv")
    class_balance = pd.read_csv(ML_TABLES / "class_balance.csv")
    holdout_diag = pd.read_csv(ML_TABLES / "split_diagnostics_holdout.csv")
    holdout_summary = pd.read_csv(ML_TABLES / "holdout_metrics_summary.csv")
    cv_summary = pd.read_csv(ML_TABLES / "cv_metrics_summary.csv")

    n_total_admissions = 58976
    n_qualifying = int(cohort_flow.loc[0, "n_rows"])
    n_patients = int(cohort_flow.loc[0, "n_patients"])
    prevalence = float(class_balance.loc[class_balance.mortality_label == 1, "proportion"].iloc[0])
    mean_train_groups = holdout_diag["n_groups_train"].mean()
    mean_test_groups = holdout_diag["n_groups_test"].mean()
    mean_pos_train = holdout_diag["positive_rate_train"].mean()
    mean_pos_test = holdout_diag["positive_rate_test"].mean()

    p_46_anchor.insert_paragraph_before("4.6. Klinik model doğrulama bulguları", style="Heading 2")
    p_46_anchor.insert_paragraph_before(
        "Tablo 12, sızıntısız erken pencere kohortunun akışını ve hasta-düzeyi bölünme "
        "tanılamasını özetler.",
        style="Normal",
    )
    build_table(
        document, p_46_anchor,
        ["Aşama", "Değer"],
        [
            ["Toplam kabul (58.976 içinden)", f"{n_total_admissions:,}".replace(",", ".")],
            ["ICU kalışı ≥48 saat (dahil edilen)",
             f"{n_qualifying:,}".replace(",", ".") + f" (%{tr(100*n_qualifying/n_total_admissions,1)})"],
            ["Benzersiz hasta (subject_id)", f"{n_patients:,}".replace(",", ".")],
            ["Mortalite prevalansı", f"%{tr(100*prevalence,1)}"],
            ["Ortalama eğitim grubu (hasta), tekrarlı ayrım", f"{tr(mean_train_groups,0)}"],
            ["Ortalama test grubu (hasta), tekrarlı ayrım", f"{tr(mean_test_groups,0)}"],
            ["Ortalama pozitif oran, eğitim", f"%{tr(100*mean_pos_train,2)}"],
            ["Ortalama pozitif oran, test", f"%{tr(100*mean_pos_test,2)}"],
        ],
    )
    p_46_anchor.insert_paragraph_before("", style="Normal")
    p_46_anchor.insert_paragraph_before(
        "Tablo 12. Sızıntısız erken pencere (48 saat) kohort akışı ve hasta-düzeyi "
        "bölünme tanılaması (n=30 tekrar ortalaması).",
        style="Caption",
    )

    p_46_anchor.insert_paragraph_before(
        "Tablo 13, tekrarlı ayrım (n=30) ve 5-katlı gruplu çapraz doğrulama "
        "sonuçlarını özetler. HistGradientBoosting en yüksek ayırt ediciliği "
        "göstermiştir (AUROC 0,814±0,007 tekrarlı ayrım, 0,814±0,011 çapraz "
        "doğrulama, AUPRC 0,402±0,016). Kalibrasyon eğimi 1,06±0,05 ile neredeyse "
        "ideal (1,0), intercept -1,68 sistematik bir kaymaya işaret eder ve "
        "olasılıkların ölçeklenmesinde ek bir yeniden kalibrasyon adımının yararlı "
        "olabileceğini düşündürür. Lojistik Regresyon daha düşük ayırt edicilikle "
        "(AUROC yaklaşık 0,77) fakat duyarlılık-özgüllük dengesinde en tutarlı "
        "sonucu üretmiştir (yaklaşık %69/%69, eşik=0,5). Rastgele Orman en yüksek "
        "AUROC'a yakın bir değer yakalasa da (0,810), varsayılan 0,5 eşiğinde çok "
        "düşük duyarlılık (yaklaşık %7,5) ile yüksek özgüllük (yaklaşık %99,3) "
        "göstermiştir. Bu, ağaç-tabanlı modelin olasılık çıktısının bu dengesiz "
        "sınıf dağılımında (mortalite %12,6) eşik-bağımlı karar noktası için yeniden "
        "kalibre edilmesi gerektiğine işaret eder ve AUROC'un eşikten bağımsız bir "
        "ayırt edicilik ölçütü olarak neden ayrı raporlandığını doğrular. Holdout ve "
        "çapraz doğrulama sonuçları birbirine çok yakındır (örneğin HistGradientBoosting "
        "AUROC'u her iki değerlendirmede de 0,8137), bu da sonucun belirli bir "
        "bölünmeye aşırı uyum göstermediğine dair ek kanıttır.",
        style="Normal",
    )

    model_labels = {"Gradient Boosting": "HistGradientBoosting",
                     "Logistic Regression": "Lojistik Regresyon",
                     "Random Forest": "Rastgele Orman"}
    rows13 = []
    for split_name, summary in [("Tekrarlı ayrım (n=30)", holdout_summary), ("5-katlı CV", cv_summary)]:
        for _, r in summary.iterrows():
            rows13.append([
                model_labels.get(r["model"], r["model"]), split_name,
                f"{tr(r['auroc_mean'],3)}±{tr(r['auroc_std'],3)}",
                f"{tr(r['auprc_mean'],3)}±{tr(r['auprc_std'],3)}",
                f"{tr(r['brier_score_mean'],3)}±{tr(r['brier_score_std'],3)}",
                f"{tr(r['calibration_intercept_mean'],2)}±{tr(r['calibration_intercept_std'],2)}",
                f"{tr(r['calibration_slope_mean'],2)}±{tr(r['calibration_slope_std'],2)}",
            ])
    build_table(
        document, p_46_anchor,
        ["Model", "Değerlendirme", "AUROC", "AUPRC", "Brier skoru",
         "Kalibrasyon intercept", "Kalibrasyon eğim"],
        rows13,
    )
    p_46_anchor.insert_paragraph_before("", style="Normal")
    p_46_anchor.insert_paragraph_before(
        "Tablo 13. Sızıntısız erken pencere kohortunda model performansı, "
        "ortalama±SS (tekrarlı ayrım n=30 ve 5-katlı gruplu çapraz doğrulama). "
        "Kalibrasyon intercept/eğim, logit(olasılık) üzerinde lojistik regresyonla "
        "hesaplanmıştır, ideal değerler sırasıyla 0 ve 1'dir.",
        style="Caption",
    )

    p_46_anchor.insert_paragraph_before(
        "Şekil 6(a-b) HistGradientBoosting, Lojistik Regresyon ve Rastgele Orman için "
        "ROC ve kesinlik-duyarlılık eğrilerini gösterir. Şekil 7 kalibrasyon "
        "(güvenilirlik) eğrisini sunar, HistGradientBoosting diyagonale en yakın "
        "modeldir. Şekil 8, Rastgele Orman'ın en önemli on beş özelliğini gösterir. "
        "Yaş, laboratuvar anormallik sayısı, laboratuvar özet istatistikleri ile "
        "solunum hızı ve sıcaklık ortalamaları öne çıkmaktadır. Hiçbir "
        "ICU-çıkışı-bağımlı değişken (yatış süresi, ICU kalış sayısı) listede yer "
        "almamaktadır, çünkü bu değişkenler bu üründe hiç hesaplanmamıştır.",
        style="Normal",
    )
    fig_roc = p_46_anchor.insert_paragraph_before("", style="Normal")
    fig_roc.add_run().add_picture(str(ML_FIGURES / "figure_ml_roc_pr_curves.png"), width=Pt(430))
    p_46_anchor.insert_paragraph_before(
        "Şekil 6. (a) ROC ve (b) kesinlik-duyarlılık eğrileri, sızıntısız erken "
        "pencere kohortunda üç model için (ayrım kümesi, tekrar 0).",
        style="Caption",
    )
    fig_calib = p_46_anchor.insert_paragraph_before("", style="Normal")
    fig_calib.add_run().add_picture(str(ML_FIGURES / "figure_ml_calibration_curve.png"), width=Pt(280))
    p_46_anchor.insert_paragraph_before(
        "Şekil 7. Kalibrasyon (güvenilirlik) eğrisi, sızıntısız erken pencere "
        "kohortunda üç model için.",
        style="Caption",
    )
    fig_imp = p_46_anchor.insert_paragraph_before("", style="Normal")
    fig_imp.add_run().add_picture(str(ML_FIGURES / "figure_ml_feature_importance.png"), width=Pt(340))
    p_46_anchor.insert_paragraph_before(
        "Şekil 8. Rastgele Orman özellik önem sıralaması, ilk 15 özellik.",
        style="Caption",
    )

    # ---------- 5.3 rewrite ----------
    p_53_b1.add_run(
        " Bu sınırlama yalnızca compact ve timeseries ETL-karşılaştırma ürünleri için "
        "geçerlidir. Bölüm 3.7'de tanımlanan ayrı, amaca özel bir erken-pencere ürünü "
        "bu sınırlamayı taşımaz."
    )
    replace_paragraph_text(
        p_53_b2,
        "Bu gereksinimler Bölüm 3.7'de tanımlanan ayrı erken-pencere ürünüyle "
        "karşılanmıştır: aynı hadm_id/subject_id'ye ait kayıtlar StratifiedGroupKFold "
        "ile aynı bölümde tutulmuş, tahmin anından (ICU girişinden 48 saat) sonra "
        "oluşan laboratuvarlar, ICU çıkış zamanı ve toplam yatış süresi dışlanmıştır. "
        "Değerlendirme AUROC, AUPRC, Brier skoru ve kalibrasyon eğrisi/intercept/"
        "slope'u içerir (Bölüm 4.6). Karar eğrisi analizi ve klinik alt grup "
        "sonuçları TRIPOD+AI'nin önerdiği tam kapsamı karşılamak üzere gelecek "
        "çalışma olarak kalmıştır [7,8,10].",
    )

    # ---------- 6.4 / 6.5 scope corrections ----------
    replace_paragraph_text(
        p_64_body,
        "Deney tek bir Windows dizüstü host, tek SSD ve tek Spark sürümüyle "
        "yürütülmüştür. Sonuçlar ayrı ağ, disk ve bellek kaynakları bulunan gerçek "
        "çok-hostlu kümeler için strong-scaling kanıtı değildir. MIMIC-III tek "
        "merkezli ve tarihsel bir kaynaktır. ETL-topolojisi karşılaştırması adalet, "
        "dış doğrulama veya hasta yararı iddiasında bulunmaz. Bölüm 4.6'daki klinik "
        "model doğrulaması da tek merkezli, dış doğrulaması yapılmamış bir "
        "kohorttadır, ayrım/eşitlik analizi ve karar eğrisi kapsam dışıdır (Bölüm "
        "6.5). `output_partitions=1` ve sekiz shuffle bölümü başka ölçeklerde farklı "
        "darboğazlar yaratabilir.",
    )
    p_65_body.add_run(
        " Bölüm 3.7'de tanımlanan ayrı erken-pencere ürünü bu paragrafın konusu "
        "değildir ve yukarıdaki sızıntı sınırlamalarını taşımaz. Buna karşın dış "
        "doğrulama, karar eğrisi analizi ve klinik alt grup performansı (örneğin "
        "yaş, cinsiyet veya tanı grubuna göre) hâlâ gelecek çalışma kapsamındadır."
    )

    # ---------- 7. Sonuç ----------
    p_conclusion.add_run(
        " Aynı ETL hattından türetilen, ICU kalışı 48 saati aşan hastalarla sınırlı, "
        "hasta-düzeyinde gruplanmış sızıntısız bir özellik uzayı, hastane-içi "
        "mortalite riskini klinik olarak makul bir ayırt edicilikle tahmin etmiştir "
        "(HistGradientBoosting AUROC 0,814±0,007, kalibrasyon eğimi 1,06±0,05). Bu, "
        "veri mühendisliği performans karakterizasyonu ile aşağı-akış model "
        "doğrulamasının aynı, tekrar üretilebilir çerçevede birlikte "
        "raporlanabileceğini göstermektedir."
    )

    # ---------- checklist ----------
    p_checklist_ml.runs[0].text = (
        "Tamamlandı: Sızıntısız erken pencere özellik seti (48 saat, ICU LOS≥48s "
        "dışlama) subject_id düzeyinde StratifiedGroupKFold ile doğrulandı. AUROC/"
        "AUPRC/Brier/kalibrasyon Bölüm 3.7/4.6'da raporlandı. Karar eğrisi analizi ve "
        "klinik alt grup performansı hâlâ gelecek çalışma kapsamındadır."
    )
    for extra_run in p_checklist_ml.runs[1:]:
        extra_run.text = ""

    # ---------- Style fix #1: reposition every table caption ABOVE its table ----------
    body = document.element.body
    children = list(body)
    for i, child in enumerate(children):
        if child.tag != qn("w:tbl"):
            continue
        # scan forward for the caption paragraph (skipping blank paragraphs in between)
        j = i + 1
        caption_el = None
        blanks = []
        while j < len(children):
            sib = children[j]
            if sib.tag != qn("w:p"):
                break
            style_el = sib.find(qn("w:pPr") + "/" + qn("w:pStyle"))
            style_val = style_el.get(qn("w:val")) if style_el is not None else None
            para_text = "".join(t.text or "" for t in sib.iter(qn("w:t")))
            if style_val and "Caption" in style_val:
                caption_el = sib
                break
            if para_text.strip() == "":
                blanks.append(sib)
                j += 1
                continue
            break
        if caption_el is not None:
            for b in blanks:
                body.remove(b)
            body.remove(caption_el)
            child.addprevious(caption_el)

    document.save(DST)

    # ---------- Style fix #2: remove semicolons throughout (period + capitalize next) ----------
    # Done as a second pass on the SAVED file's paragraph text, after all structural
    # inserts/moves, so every new and pre-existing paragraph is covered uniformly.
    document2 = Document(DST)

    def desemicolon(text: str) -> str:
        def repl(m: re.Match) -> str:
            before, after_char = m.group(1), m.group(2)
            return f"{before}. {after_char.upper()}"
        # "; " or ";" followed by a letter -> ". " + capitalize next letter
        return re.sub(r"([^\s;])\s*;\s*([a-zçğıöşüA-ZÇĞİÖŞÜ0-9])", repl, text)

    changed = 0
    for p in document2.paragraphs:
        if ";" not in p.text:
            continue
        new_text = desemicolon(p.text)
        if new_text == p.text:
            continue
        changed += 1
        # Preserve the paragraph's first run formatting; collapse all runs into one,
        # matching the same technique used above for targeted rewrites.
        for run in list(p.runs)[1:]:
            run.text = ""
        if p.runs:
            p.runs[0].text = new_text
        else:
            p.add_run(new_text)
    for t in document2.tables:
        for row in t.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    if ";" not in p.text:
                        continue
                    new_text = desemicolon(p.text)
                    if new_text == p.text:
                        continue
                    changed += 1
                    for run in list(p.runs)[1:]:
                        run.text = ""
                    if p.runs:
                        p.runs[0].text = new_text
                    else:
                        p.add_run(new_text)

    document2.save(DST)
    print(f"Saved: {DST}")
    print(f"Semicolon sweep: {changed} paragraphs/cells changed")


if __name__ == "__main__":
    main()
