from __future__ import annotations

from pathlib import Path

from docx import Document

DOCX = Path(__file__).resolve().parent / "MIMICIII_Spark_ETL_Makale_Literatur_Temelli_v10.docx"


def main() -> None:
    document = Document(DOCX)
    paras = document.paragraphs

    p_intro_link_anchor = paras[14]   # §1 body para 5 -> insert new linking para before it
    p_25_anchor = paras[41]           # heading "3. Materyal ve Yöntem" -> insert §2.5 before it
    p_24_synth = paras[40]            # §2.4 synthesis paragraph -> append pointer to §2.5
    p_ref_anchor = paras[172]         # heading "Gönderim öncesi..." -> insert new refs before it

    # ---------- Short linking paragraph in Introduction (§1) ----------
    p_intro_link_anchor.insert_paragraph_before(
        "Ayrı bir güncel çizgi, aynı MIMIC ailesi veritabanları üzerinde ICU mortalite "
        "tahmininde sızıntı kontrolü, hasta-düzeyinde bölünme ve TRIPOD+AI uyumlu "
        "raporlamaya artan bir vurgudur (Bölüm 2.5). Bu çizgi veri hazırlama hattının "
        "kendisinden çok, üretilen özellik uzayının aşağı-akış model doğrulamasında "
        "nasıl kullanıldığıyla ilgilidir ve mevcut çalışmanın Bölüm 3.7/4.6'daki "
        "sızıntısız erken pencere katkısıyla doğrudan kesişir.",
        style="Normal",
    )

    # ---------- §2.4 synthesis: pointer to §2.5 ----------
    p_24_synth.add_run(
        " Bölüm 2.5, bu sentezi ICU mortalite tahmininde sızıntı kontrolü ve kalibrasyon "
        "raporlamasına odaklanan 2023-2026 dönemi çalışmalarıyla genişletir."
    )

    # ---------- New §2.5 ----------
    p_25_anchor.insert_paragraph_before(
        "2.5. Sızıntı kontrolü ve kalibrasyon odaklı güncel ICU mortalite tahmini "
        "çalışmaları", style="Heading 2",
    )
    p_25_anchor.insert_paragraph_before(
        "BMC Medical Informatics and Decision Making'de 2023-2026 arasında yayımlanan "
        "çalışmalar, MIMIC tabanlı ICU mortalite tahmininde sızıntı kontrolü, "
        "kalibrasyon ve raporlama standardına artan bir vurguyu göstermektedir. "
        "Bouvarel ve ark. [30] MIMIC-III üzerinde 46.520 hastadan başlayarak ICU "
        "kalışı 48 saatten kısa olan hastaları dışlamış (mevcut çalışmanın kullandığı "
        "aynı eşik) ve 36 saatlik gözlem penceresinden 12-24 saat sonrasını tahmin "
        "etmiştir. Tam-vaka analizinde hastaların %75'i dışlanırken çoklu imputasyon "
        "hastaların %95'ini (16.532/17.373) koruyarak 0,777-0,789 AUC aralığında "
        "karşılaştırılabilir ayırt edicilik elde etmiştir. Bu çalışma, eksik veri "
        "stratejisinin kohort büyüklüğü üzerindeki etkisini nicel olarak göstermesi "
        "bakımından, mevcut çalışmanın eksik-veri göstergesi (missingness indicator) "
        "yaklaşımıyla doğrudan karşılaştırılabilir bir emsaldir.",
        style="Normal",
    )
    p_25_anchor.insert_paragraph_before(
        "Yeh ve ark. [31], MIMIC-IV ve bir Taipei hastanesinden 79.657 kabul "
        "kullanarak ICU kabulü anında (AUROC 0,856) ve kabulden 24 saat sonra "
        "(AUROC 0,910) olmak üzere iki ayrı erken tahmin zaman noktası "
        "karşılaştırmıştır. Bu, tahmin ufku uzadıkça ayırt ediciliğin arttığını nicel "
        "olarak göstermektedir ve mevcut çalışmanın 48 saatlik pencere seçiminin "
        "AUROC-bilgi ödünleşimi bağlamında nereye düştüğünü yorumlamaya yardımcı "
        "olur. Samadi ve ark. [32], RWTH Aachen ICU verisiyle TRIPOD+AI kılavuzuna "
        "uygun raporlama yapmış ve GPT-4o destekli hibrit XGBoost modeliyle "
        "0,780-0,784 AUC-ROC elde etmiştir. Bu çalışma, derginin TRIPOD+AI uyumlu "
        "raporlamayı güncel bir beklenti olarak desteklediğini göstermektedir. Shetty "
        "ve ark. [33], pediatrik solunum hastalığı sınıflandırmasında sızıntı-kontrollü "
        "ve SHAP tabanlı bir çerçeve önererek veri sızıntısının klinik ML "
        "güvenilirliğini nasıl tehdit ettiğini vurgulamıştır. Bu, mevcut çalışmanın "
        "erken-pencere ürününün tasarım motivasyonuyla doğrudan örtüşmektedir.",
        style="Normal",
    )
    p_25_anchor.insert_paragraph_before(
        "Buna karşın aynı dönemde MIMIC-III sepsis-3 mortalite tahmininde bildirilen "
        "AUC değerleri geniş bir aralığa yayılmaktadır: Yu ve ark. [34] LightGBM ile "
        "0,983 AUC, Rahman ve ark. [35] ise yığınlama (stacking) meta-sınıflandırıcıyla "
        "%95,52 doğruluk bildirmiştir. Bu denli yüksek ayırt edicilik değerleri, "
        "alanda özellik tanımı ve veri bölme stratejisindeki farklılıkların sonuçları "
        "ne ölçüde etkileyebileceğini göstermektedir. Mevcut çalışma, hasta-düzeyinde "
        "gruplanmış bölünme ve açıkça tanımlanmış 48 saatlik tahmin ufkuyla, bu geniş "
        "aralık içinde klinik olarak temkinli ve yöntemsel olarak sınanabilir bir "
        "konum almaktadır (AUROC 0,814±0,007). Bu alt bölümde taranan çalışmaların "
        "hiçbiri Spark yürütme topolojisini veya kaynak-kısıtlı tek-host ETL "
        "performansını incelememektedir; bu, Bölüm 2.3-2.4'te tanımlanan sistem "
        "boşluğunun bu dergide de güncel biçimde açık kaldığını doğrulamaktadır.",
        style="Normal",
    )

    # ---------- New references (30-35) ----------
    new_refs = [
        "30. Bouvarel B, Carrat F, Lapidus N. Updating mortality risk estimation in "
        "intensive care units from high-dimensional electronic health records with "
        "incomplete data. BMC Medical Informatics and Decision Making. 2023;23:170. "
        "doi:10.1186/s12911-023-02264-7.",
        "31. Yeh YC, Kuo YT, Kuo KC, et al. Early prediction of mortality upon "
        "intensive care unit admission. BMC Medical Informatics and Decision Making. "
        "2024. doi:10.1186/s12911-024-02807-6.",
        "32. Samadi ME, Nikulina K, Fritsch SJ, Schuppert A. GPT-4o and the quest for "
        "machine learning interpretability in ICU risk of death prediction. BMC "
        "Medical Informatics and Decision Making. 2025. doi:10.1186/s12911-025-03224-z.",
        "33. Shetty AP, Shetty S, Hegde P, Shetty S, Shetty N. A leakage-controlled "
        "and SHAP driven machine learning framework for paediatric respiratory "
        "disease classification using Indian hospital EHR data. BMC Medical "
        "Informatics and Decision Making. 2026. doi:10.1186/s12911-026-03493-2.",
        "34. Yu Z, Ashrafi N, Li H, Alaei K, Pishgar M. Prediction of 30-day "
        "mortality for ICU patients with Sepsis-3. BMC Medical Informatics and "
        "Decision Making. 2024. doi:10.1186/s12911-024-02629-6.",
        "35. Rahman MS, Islam KR, Prithula J, et al. Machine learning-based "
        "prognostic model for 30-day mortality prediction in Sepsis-3. BMC Medical "
        "Informatics and Decision Making. 2024. doi:10.1186/s12911-024-02655-4.",
    ]
    for ref in new_refs:
        p_ref_anchor.insert_paragraph_before(ref, style="Normal")

    document.save(DOCX)
    print("Saved literature update to", DOCX)


if __name__ == "__main__":
    main()
