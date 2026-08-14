from __future__ import annotations

from pathlib import Path

from docx import Document

DOCX = Path(__file__).resolve().parent / "MIMICIII_Spark_ETL_Makale_Literatur_Temelli_v10.docx"


def replace_paragraph_text(paragraph, new_text: str) -> None:
    for run in list(paragraph.runs)[1:]:
        run.text = ""
    if paragraph.runs:
        paragraph.runs[0].text = new_text
    else:
        paragraph.add_run(new_text)


def main() -> None:
    document = Document(DOCX)
    paras = document.paragraphs

    replace_paragraph_text(
        paras[1],
        "Esma Fazilet Karagülle¹ (sorumlu yazar), Esra Odabaş Yıldırım²  |  "
        "¹Atatürk Üniversitesi, Bilgisayar Mühendisliği, "
        "esmafazilet.karagulle17@ogr.atauni.edu.tr  |  "
        "²Atatürk Üniversitesi, Yazılım Mühendisliği, esra.odabas@atauni.edu.tr",
    )
    replace_paragraph_text(paras[138], "Çıkar çatışması: Yazarlar çıkar çatışması bildirmemektedir.")
    replace_paragraph_text(paras[139], "Finansman: Bu çalışma özel bir fon almamıştır.")
    replace_paragraph_text(
        paras[140],
        "Yazar katkıları: Esma Fazilet Karagülle — Kavramsallaştırma, Metodoloji, "
        "Yazılım, Doğrulama, Formal analiz, Araştırma, Veri küratörlüğü, Yazma "
        "(orijinal taslak), Görselleştirme. Esra Odabaş Yıldırım — Kavramsallaştırma, "
        "Metodoloji, Danışmanlık (süpervizyon), Yazma (inceleme ve düzenleme).",
    )
    replace_paragraph_text(
        paras[141],
        "Kod ve toplulaştırılmış sonuçların erişilebilirliği: Kod deposu "
        "https://github.com/esma6/icu_mortality_project adresinde paylaşılacaktır "
        "(depo hazırlığı ve README tamamlandıktan sonra bağlantı doğrulanacaktır). "
        "Paylaşım paketi hasta düzeyi veri içermez. Kaynak kod, anonimleştirilmiş "
        "timing/resource günlükleri ve istatistik üretim betiklerini içerir.",
    )
    replace_paragraph_text(
        paras[173],
        "Tamamlandı: Yazar adı, kurum ve e-posta dolduruldu (sorumlu yazar: Esma "
        "Fazilet Karagülle). Kod deposu bağlantısı GitHub'a yükleme tamamlandıktan "
        "sonra doğrulanacaktır.",
    )

    document.save(DOCX)
    print("Saved author info updates to", DOCX)


if __name__ == "__main__":
    main()
